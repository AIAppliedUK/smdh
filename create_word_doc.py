#!/usr/bin/env python3
"""
Generate comprehensive SMDH Architecture Design Options Word document
Uses python-docx library to create a professional document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Create document
doc = Document()

# Set up styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Title page
title = doc.add_heading('Smart Manufacturing Data Hub (SMDH)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Architecture Design Options')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_format = subtitle.runs[0]
subtitle_format.font.size = Pt(18)
subtitle_format.font.bold = True
subtitle_format.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()
doc.add_paragraph()

# Document information table
info_table = doc.add_table(rows=5, cols=2)
info_table.style = 'Light Grid Accent 1'
info_cells = info_table.rows[0].cells
info_cells[0].text = 'Version'
info_cells[1].text = '2.0 FINAL'
info_cells = info_table.rows[1].cells
info_cells[0].text = 'Date'
info_cells[1].text = 'October 2025'
info_cells = info_table.rows[2].cells
info_cells[0].text = 'Status'
info_cells[1].text = 'For Review'
info_cells = info_table.rows[3].cells
info_cells[0].text = 'Classification'
info_cells[1].text = 'Internal Use'
info_cells = info_table.rows[4].cells
info_cells[0].text = 'Author'
info_cells[1].text = 'AI Applied Technical Team'

doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', 1)
toc_items = [
    '1. Executive Summary',
    '2. Document Purpose and Audience',
    '3. System Overview and Requirements',
    '4. Architecture Option A: Flink-Based AWS-Heavy Architecture',
    '   4.1 Overview',
    '   4.2 Architecture Diagram',
    '   4.3 Detailed Component Description',
    '   4.4 Data Flow and Processing',
    '   4.5 Strengths and Limitations',
    '5. Architecture Option B: Snowflake-Leveraged Architecture',
    '   5.1 Overview',
    '   5.2 Architecture Diagram',
    '   5.3 Detailed Component Description',
    '   5.4 Data Flow and Processing',
    '   5.5 Strengths and Limitations',
    '6. Architecture Option C: AWS IoT SiteWise-Based Architecture',
    '   6.1 Overview',
    '   6.2 Architecture Diagram',
    '   6.3 Detailed Component Description',
    '   6.4 Data Flow and Processing',
    '   6.5 Strengths and Limitations',
    '7. Architecture Option D: Pure AWS-Native Architecture (Timestream + Grafana)',
    '   7.1 Overview',
    '   7.2 Architecture Diagram',
    '   7.3 Detailed Component Description',
    '   7.4 Data Flow and Processing',
    '   7.5 Strengths and Limitations',
    '8. Architecture Comparison',
    '   8.1 Cost Comparison',
    '   8.2 Complexity and Skills Required',
    '   8.3 Performance and Latency',
    '   8.4 Multi-Tenancy Approach',
    '   8.5 Decision Framework',
    '9. Recommendations',
    '10. Glossary of Terms',
]

for item in toc_items:
    doc.add_paragraph(item, style='List Bullet' if item.startswith('   ') else 'List Number')

doc.add_page_break()

print("Document structure created successfully")
print("Continuing with content sections...")

# Save
doc.save('/Users/david/projects/smdh/docs/architecture/SMDH-Infrastructure-Design-Options-v2.docx')
print("Document saved successfully!")

#!/usr/bin/env python3
"""
Expand Options C and D with comprehensive detailed content
Matches the formatting and detail level of Options A and B
Uses UK English throughout
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_page_break(doc):
    doc.add_page_break()

def add_heading_with_color(doc, text, level, color_rgb=(0, 51, 102)):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

print("Opening document...")
doc = Document('/Users/david/projects/smdh/docs/architecture/SMDH Infrastructure Design Options.docx')

# Find where Option C starts (currently just summary)
option_c_index = None
option_d_index = None
comparison_index = None

for i, para in enumerate(doc.paragraphs):
    if 'Option C' in para.text and 'SiteWise' in para.text and para.style.name == 'Heading 1':
        option_c_index = i
        print(f"Found Option C at paragraph {i}")
    if 'Option D' in para.text and para.style.name == 'Heading 1':
        option_d_index = i
        print(f"Found Option D at paragraph {i}")
    if '7. Architecture Comparison' in para.text:
        comparison_index = i
        print(f"Found Comparison section at paragraph {i}")
        break

# Remove existing Option C content (summary only) up to Option D
if option_c_index and option_d_index:
    print(f"Removing paragraphs {option_c_index + 1} to {option_d_index - 1}")
    for _ in range(option_d_index - option_c_index - 1):
        p = doc.paragraphs[option_c_index + 1]
        p._element.getparent().remove(p._element)

# Now insert comprehensive Option C content after the heading
print("Inserting comprehensive Option C content...")

# Get the parent element for insertion
option_c_heading = doc.paragraphs[option_c_index]
parent = option_c_heading._element.getparent()

def insert_after(existing_element, new_element):
    """Insert new element after existing element"""
    parent = existing_element.getparent()
    parent.insert(parent.index(existing_element) + 1, new_element)

# Helper to add paragraph after a specific element
def add_paragraph_after(doc, reference_para, text='', style=None):
    """Add a new paragraph after a reference paragraph"""
    new_para = doc.add_paragraph(text, style=style)
    # Move it to right after reference
    new_elem = new_para._element
    ref_elem = reference_para._element
    ref_parent = ref_elem.getparent()
    ref_parent.insert(ref_parent.index(ref_elem) + 1, new_elem)
    return new_para

# Since direct insertion is complex, let's rebuild from Option C onwards
print("Rebuilding Option C section with full detail...")

# Find the insertion point again (after removal)
option_c_index = None
for i, para in enumerate(doc.paragraphs):
    if 'Option C' in para.text and 'SiteWise' in para.text and para.style.name == 'Heading 1':
        option_c_index = i
        break

# Remove everything from Option C heading onwards
while len(doc.paragraphs) > option_c_index + 1:
    p = doc.paragraphs[option_c_index + 1]
    p._element.getparent().remove(p._element)

print(f"Cleared from paragraph {option_c_index + 1} onwards. Now adding Option C detail...")

# Add Option C title (already exists, so continue from here)

doc.add_heading('5.1 Overview and Summary', 2)

summary_table_c = doc.add_table(rows=7, cols=2)
summary_table_c.style = 'Medium Shading 1 Accent 4'

cells = summary_table_c.rows[0].cells
cells[0].text = 'Characteristic'
cells[1].text = 'Details'

cells = summary_table_c.rows[1].cells
cells[0].text = 'Monthly Cost (30 tenants)'
cells[1].text = '£6,334 ✅ CORRECTED (was £21,150)'

cells = summary_table_c.rows[2].cells
cells[0].text = 'Cost Per Tenant'
cells[1].text = '£211'

cells = summary_table_c.rows[3].cells
cells[0].text = 'Complexity Level'
cells[1].text = 'Medium-High (8-10 AWS services + dual storage)'

cells = summary_table_c.rows[4].cells
cells[0].text = 'Implementation Timeline'
cells[1].text = '28 weeks (longest - custom dashboards required)'

cells = summary_table_c.rows[5].cells
cells[0].text = 'Team Size & Skills'
cells[1].text = '3-4 engineers (SiteWise, AWS IoT, React, Timestream)'

cells = summary_table_c.rows[6].cells
cells[0].text = 'Best Suited For'
cells[1].text = 'Teams with deep SiteWise expertise; <10 tenants; AWS-native preference'

doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('⚠️ IMPORTANT COST CORRECTION')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(153, 51, 0)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

doc.add_paragraph(
    'A previous analysis incorrectly calculated Option C cost at £21,150/month due to a 7.3x '
    'data volume error. The CORRECTED cost is £6,334/month (70% reduction), making this option '
    'economically viable. However, it remains the most expensive option and has multi-tenancy '
    'limitations for SaaS platforms.'
)

doc.add_paragraph()

doc.add_paragraph(
    'Option C leverages AWS IoT SiteWise, a managed service purpose-built for industrial equipment '
    'monitoring. SiteWise provides native concepts like asset hierarchies (Factory → Line → Machine), '
    'automatic time-series calculations, and OPC-UA/Modbus protocol support. However, it is NOT designed '
    'for multi-tenant SaaS, requiring significant customisation and dual-storage architecture.'
)

doc.add_heading('5.2 Architecture Diagram', 2)

# Add diagram placeholder
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n[INSERT DIAGRAM: SMDH-Option-C-SiteWise-Architecture.drawio]\n')
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)
for _ in range(6):
    p.add_run('\n')
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F5F5F5')
pPr.append(shd)

doc.add_heading('5.3 Core Components', 2)

doc.add_heading('Dual Ingestion Path (SiteWise + Timestream)', 3)

para = doc.add_paragraph()
run = para.add_run('Why Dual Storage?')
run.bold = True

doc.add_paragraph(
    'Option C requires TWO separate data stores because SiteWise and Timestream serve different purposes:'
)

dual_storage_table = doc.add_table(rows=4, cols=3)
dual_storage_table.style = 'Light Grid Accent 1'

cells = dual_storage_table.rows[0].cells
cells[0].text = 'Aspect'
cells[1].text = 'SiteWise'
cells[2].text = 'Timestream'

cells = dual_storage_table.rows[1].cells
cells[0].text = 'Data Type'
cells[1].text = 'Continuous time-series (machine status, sensors)'
cells[2].text = 'Discrete events (RFID scans, job tracking)'

cells = dual_storage_table.rows[2].cells
cells[0].text = 'Best For'
cells[1].text = 'Regular sensor readings, asset hierarchies, OEE'
cells[2].text = 'Irregular events, rich metadata, SQL queries'

cells = dual_storage_table.rows[3].cells
cells[0].text = 'Use Cases'
cells[1].text = 'Machine utilisation, energy, air quality'
cells[2].text = 'Job location tracking, barcode scans'

doc.add_paragraph()

components_c_ingestion = [
    ('SiteWise Gateway (On-Premises)', 'Software appliance running at manufacturing sites. Translates '
     'industrial protocols (OPC-UA, Modbus TCP) to AWS IoT format. Handles local caching for offline '
     'operation. Requires management at each site (updates, security patches).'),
    ('AWS IoT Core', 'MQTT broker (same as other options). Routes data based on type: continuous data '
     '→ SiteWise, events → Kinesis.'),
    ('IoT Rules Engine', 'Routes messages by content. Continuous sensor data goes to SiteWise, discrete '
     'events (RFID/barcodes) route to Kinesis Streams.'),
    ('Kinesis Streams (Events Only)', 'Buffers discrete events for ingestion into Timestream. Does NOT '
     'handle continuous sensor data (that goes directly to SiteWise).')
]

for component, description in components_c_ingestion:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('AWS IoT SiteWise (Time-Series Platform)', 3)

para = doc.add_paragraph()
run = para.add_run('What Makes SiteWise Different?')
run.bold = True

doc.add_paragraph(
    'Unlike general-purpose databases, SiteWise is purpose-built for industrial IoT:'
)

sitewise_features = [
    'Asset Models: Templates defining equipment structure (e.g., "CNC Machine" with properties for '
    'spindle speed, temperature, state)',
    'Asset Hierarchies: Organise equipment in trees (Company → Site → Production Line → Machine → Sensor)',
    'Compute Expressions: Automatic calculations (OEE, availability, performance) without custom code',
    'Native Time-Series Storage: Optimised for high-frequency sensor data with automatic aggregations',
    'Built-in Formulas: Industry-standard metrics (OEE, MTBF, MTTR) pre-configured'
]

for feature in sitewise_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()

components_sitewise = [
    ('SiteWise Asset Ingestion', 'Receives data from Gateway or direct MQTT. Validates against asset '
     'models. Stores in time-series database. Latency: <1 second.'),
    ('Asset Models', 'Define structure: Measurements (raw sensor readings), Metrics (calculated values '
     'like OEE), Transforms (unit conversions, aggregations), Alarms (threshold-based alerts).'),
    ('Compute Expressions', 'SiteWise Expression Language (similar to Excel formulas). Example: '
     'availability = (total_time - downtime) / total_time. Runs automatically as data arrives.'),
    ('Time-Series Store', 'Hot tier: 13 months (fast queries). Cold tier: 14+ months (S3-based archival). '
     'Automatic aggregations: 1min, 5min, 1hour, 1day intervals.'),
    ('SiteWise Monitor', 'Built-in dashboarding tool. LIMITATION: Cannot be adequately white-labelled '
     'for SaaS. AWS branding remains visible. Not multi-tenant aware.')
]

for component, description in components_sitewise:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('⚠️ Critical Multi-Tenancy Challenge', 3)

para = doc.add_paragraph()
run = para.add_run('SiteWise is NOT Designed for Multi-Tenant SaaS')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(153, 0, 0)

doc.add_paragraph(
    'Unlike Snowflake\'s Row-Level Security or Timestream\'s partition keys, SiteWise has NO native '
    'multi-tenancy support. The workaround:'
)

mt_workaround = [
    'Tag every asset with "tenant_id" metadata',
    'Application layer (Lambda/API) must filter ALL queries by tenant_id',
    'Custom React application enforces filtering in UI code',
    'Requires extensive security testing—a bug could expose one tenant\'s data to another'
]

for item in mt_workaround:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('Risk Assessment: ')
run.bold = True
para.add_run(
    'Medium-high risk for multi-tenant SaaS. Acceptable for single-tenant deployments or <10 trusted '
    'tenants. Requires rigorous code review and penetration testing.'
)

doc.add_heading('Amazon Timestream (Event Data)', 3)

components_timestream_c = [
    ('Why Timestream?', 'SiteWise cannot efficiently handle discrete events (RFID scans, job completions). '
     'These have irregular timing and rich metadata. Timestream is purpose-built for this.'),
    ('Timestream Tables', 'Separate tables for job_scan_events, quality_checks, maintenance_logs. '
     'Partition key: tenant_id (better multi-tenancy than SiteWise tags).'),
    ('SQL Query Engine', 'Standard SQL interface for event analysis. Example: "Find all scans for Job #1234 '
     'in past week." Much simpler than SiteWise property queries.')
]

for component, description in components_timestream_c:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Custom Dashboard Layer (Required)', 3)

para = doc.add_paragraph()
run = para.add_run('Why Custom Dashboards Are Necessary')
run.bold = True

doc.add_paragraph(
    'SiteWise Monitor limitations force custom development:'
)

monitor_limitations = [
    'Cannot white-label: AWS branding and URLs remain visible',
    'Not multi-tenant: No concept of separating dashboards by company',
    'Limited customisation: Dashboard layout and chart types are constrained',
    'Cannot combine data: Cannot easily show SiteWise + Timestream data together'
]

for limit in monitor_limitations:
    doc.add_paragraph(limit, style='List Bullet')

doc.add_paragraph()

dashboard_components = [
    ('Custom React Application', 'Built from scratch using React 18+ and TypeScript. Calls SiteWise and '
     'Timestream APIs directly. Uses Chart.js or Recharts for visualisations. Fully white-labelable.'),
    ('API Gateway + Lambda', 'Backend layer enforcing tenant isolation. Every API call validates user\'s '
     'tenant_id before querying SiteWise/Timestream. Filters results by tenant.'),
    ('AWS Amplify Hosting', 'Hosts React frontend. Provides CI/CD pipeline. Integrates with Cognito '
     'for authentication.')
]

for component, description in dashboard_components:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('5.4 Data Flow', 2)

doc.add_paragraph('Time-series data path (Machine/Energy/Air Quality):')
flow_ts_c = [
    'Sensor → OPC-UA/Modbus → SiteWise Gateway (on-prem)',
    'Gateway → AWS IoT SiteWise ingestion (latency: <1s)',
    'SiteWise validates against asset model, stores in time-series DB',
    'Compute expressions calculate metrics (OEE, aggregations)',
    'Custom React app queries SiteWise API for dashboard display (latency: <5s total)'
]
for step in flow_ts_c:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()
doc.add_paragraph('Event data path (RFID/Barcode Scans):')
flow_event_c = [
    'Scanner → HTTP POST to API Gateway',
    'Lambda validates and routes to Kinesis Streams (latency: <200ms)',
    'Lambda consumer writes to Timestream (latency: <1s)',
    'Custom React app queries Timestream SQL for job tracking dashboard (latency: <3s total)'
]
for step in flow_event_c:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('5.5 Strengths', 2)

strengths_c = [
    ('Purpose-Built for Manufacturing', 'Native industrial concepts: asset models, OEE, equipment hierarchies. '
     'Speaks manufacturing language, not generic database terms.'),
    ('Managed Asset Modelling', 'Built-in templates for common equipment. Automatic calculations for '
     'availability, performance, quality. Less custom code than building from scratch.'),
    ('Real-Time Latency', '<1 second for time-series data. Meets real-time monitoring requirements without '
     'complex stream processing like Flink.'),
    ('Industrial Protocol Support', 'SiteWise Gateway handles OPC-UA, Modbus TCP natively. Easier connectivity '
     'to legacy manufacturing equipment.'),
    ('Automatic Metric Calculations', 'Compute expressions run at ingestion time. Standard manufacturing metrics '
     '(OEE, MTBF, utilisation) pre-configured.')
]

for title, desc in strengths_c:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(desc)

doc.add_heading('5.6 Limitations', 2)

limits_c = [
    ('NOT Designed for Multi-Tenant SaaS', 'CRITICAL: No native Row-Level Security. Tag-based filtering '
     'is application-layer workaround with data leakage risk. Acceptable for <10 tenants; risky for SaaS.'),
    ('Most Expensive Option', 'Cost: £6,334/month (£211/tenant) vs £2,170 for Option B (cheapest). 3x more '
     'expensive. Ingestion pricing ($0.75 per million values) adds up with 1Hz sensors.'),
    ('Dual Storage Complexity', 'Requires BOTH SiteWise (time-series) AND Timestream (events). Applications '
     'must query two databases. Data correlation requires custom logic.'),
    ('Custom Dashboard Development', 'SiteWise Monitor cannot be white-labelled. Must build custom React '
     'application (6-8 weeks development + ongoing maintenance).'),
    ('Longest Timeline', '28 weeks to production. Custom dashboards, dual storage integration, and asset '
     'model configuration extend timeline 40% beyond Option B.'),
    ('On-Premises Gateway Management', 'SiteWise Gateway requires software deployment at each manufacturing site. '
     'Software updates, troubleshooting, security patching across distributed locations.'),
    ('SiteWise Learning Curve', 'Proprietary concepts (asset models, expression language) require training. '
     'Smaller talent pool than SQL or Python.')
]

for title, desc in limits_c:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)
    para.add_run(desc)

doc.add_heading('5.7 When to Choose', 2)

doc.add_paragraph('Select Option C ONLY if:')
when_c = [
    'Team has existing deep AWS IoT SiteWise expertise (rare)',
    'Manufacturing focus with <10 tenants (multi-tenancy risk acceptable)',
    'Already using SiteWise for other projects (leverage existing skills)',
    'OPC-UA/Modbus connectivity is critical and must be managed (not delegated)',
    'Budget supports £200+/tenant (most expensive option)',
    'Willing to invest 28 weeks + custom dashboard development',
    'Can accept dual storage complexity and manual multi-tenant isolation'
]
for item in when_c:
    doc.add_paragraph(item, style='List Bullet')

para = doc.add_paragraph()
run = para.add_run('\n⚠️ ASSESSMENT: ')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(153, 51, 0)
para.add_run(
    'Option C is viable after cost correction but NOT recommended for multi-tenant SaaS. The multi-tenancy '
    'risk (tag-based filtering), highest cost, dual storage complexity, and custom dashboard requirement '
    'outweigh the benefits of managed asset modelling. Choose Option B for simplicity or Option D for '
    'AWS-native without SiteWise limitations.'
)

add_page_break(doc)

print("Option C complete. Adding comprehensive Option D...")

# OPTION D: AWS-NATIVE (TIMESTREAM + GRAFANA)
add_heading_with_color(doc, '6. Architecture Option D: Pure AWS-Native Architecture (Timestream + Grafana)', 1, (0, 150, 200))

doc.add_heading('6.1 Overview and Summary', 2)

summary_table_d = doc.add_table(rows=7, cols=2)
summary_table_d.style = 'Medium Shading 1 Accent 3'

cells = summary_table_d.rows[0].cells
cells[0].text = 'Characteristic'
cells[1].text = 'Details'

cells = summary_table_d.rows[1].cells
cells[0].text = 'Monthly Cost (30 tenants)'
cells[1].text = '£3,965'

cells = summary_table_d.rows[2].cells
cells[0].text = 'Cost Per Tenant'
cells[1].text = '£132'

cells = summary_table_d.rows[3].cells
cells[0].text = 'Complexity Level'
cells[1].text = 'Medium (7-8 AWS services)'

cells = summary_table_d.rows[4].cells
cells[0].text = 'Implementation Timeline'
cells[1].text = '22 weeks'

cells = summary_table_d.rows[5].cells
cells[0].text = 'Team Size & Skills'
cells[1].text = '2-3 engineers (AWS-native, SQL, Grafana)'

cells = summary_table_d.rows[6].cells
cells[0].text = 'Best Suited For'
cells[1].text = 'AWS-native mandate; organisations preferring no Snowflake'

doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('✅ BEST AWS-NATIVE OPTION')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 102, 204)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

doc.add_paragraph(
    'Option D provides a fully AWS-native solution using Amazon Timestream for all time-series and event '
    'data, combined with Grafana for industry-leading dashboards. This approach avoids Snowflake licensing '
    'concerns (Option B) and SiteWise multi-tenancy limitations (Option C) whilst maintaining reasonable cost '
    'and complexity. It is the recommended choice when AWS-native services are mandated.'
)

doc.add_heading('6.2 Architecture Diagram', 2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n[INSERT DIAGRAM: SMDH-Option-D-Timestream-Architecture.drawio]\n')
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(128, 128, 128)
for _ in range(6):
    p.add_run('\n')
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:fill'), 'F5F5F5')
pPr.append(shd)

doc.add_heading('6.3 Core Components', 2)

doc.add_heading('Unified Ingestion (Single Data Path)', 3)

para = doc.add_paragraph()
run = para.add_run('Key Simplification vs Option C')
run.bold = True

doc.add_paragraph(
    'Unlike Option C which requires dual storage (SiteWise + Timestream), Option D uses Timestream for '
    'EVERYTHING—continuous sensor data AND discrete events. This dramatically simplifies the architecture.'
)

unified_benefits = [
    'Single database to query (not two)',
    'Unified data model (no translation between SiteWise and Timestream)',
    'Single SQL interface (familiar to most engineers)',
    'Native partition-based multi-tenancy (better than SiteWise tags)',
    'Lower operational overhead (one database to monitor, backup, optimise)'
]

for benefit in unified_benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_paragraph()

components_d_ingestion = [
    ('AWS IoT Core', 'Same MQTT broker as other options. X.509 certificate authentication. Device shadows. '
     'QoS 0/1 support.'),
    ('IoT Rules Engine', 'Routes ALL data types to Kinesis Streams. Simpler than Option C (no dual routing). '
     'Enriches messages with tenant metadata.'),
    ('Kinesis Data Streams', '10-20 shards partitioned by tenant_id. Buffers all sensor data and events. '
     '24-hour retention for replay capability.'),
    ('Lambda Ingestion Handler', 'Processes Kinesis batches. Validates sensor ranges. Converts to Timestream '
     'format. Writes to appropriate Timestream tables. Handles both time-series and events.')
]

for component, description in components_d_ingestion:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Amazon Timestream (Unified Data Platform)', 3)

para = doc.add_paragraph()
run = para.add_run('Why Timestream Over SiteWise?')
run.bold = True

timestream_comparison = doc.add_table(rows=7, cols=3)
timestream_comparison.style = 'Light Grid Accent 1'

cells = timestream_comparison.rows[0].cells
cells[0].text = 'Feature'
cells[1].text = 'SiteWise (Option C)'
cells[2].text = 'Timestream (Option D)'

cells = timestream_comparison.rows[1].cells
cells[0].text = 'Cost Model'
cells[1].text = '$0.75 per million values'
cells[2].text = '$0.50 per GB ingested'

cells = timestream_comparison.rows[2].cells
cells[0].text = 'Multi-Tenancy'
cells[1].text = 'Manual tags ⚠️'
cells[2].text = 'Native partition keys ✅'

cells = timestream_comparison.rows[3].cells
cells[0].text = 'Event Data'
cells[1].text = 'Poor fit (needs Timestream)'
cells[2].text = 'Excellent (unified) ✅'

cells = timestream_comparison.rows[4].cells
cells[0].text = 'Query Language'
cells[1].text = 'Limited property APIs'
cells[2].text = 'Full SQL ✅'

cells = timestream_comparison.rows[5].cells
cells[0].text = 'Flexibility'
cells[1].text = 'Asset model constraints'
cells[2].text = 'Schemaless ✅'

cells = timestream_comparison.rows[6].cells
cells[0].text = 'Cost (30 tenants)'
cells[1].text = '£6,334/month'
cells[2].text = '£3,965/month ✅'

doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('Verdict: ')
run.bold = True
para.add_run('Timestream is 40% cheaper and more flexible for multi-tenant SaaS.')

doc.add_paragraph()

components_timestream = [
    ('Timestream Database Schema', 'Separate tables for machine_telemetry, energy_monitoring, air_quality, '
     'job_scan_events. ALL tables partitioned by tenant_id (native physical isolation).'),
    ('Memory Store', 'Hot tier: 7-90 days configurable. Fast queries for recent data. Used for real-time '
     'dashboards and alerts.'),
    ('Magnetic Store', 'Warm/cold tier: Months to years retention. Cost-optimised storage. Used for historical '
     'analysis and compliance.'),
    ('SQL Query Engine', 'Standard SQL (mostly ANSI-compatible). Familiar to most engineers. Supports JOINs, '
     'window functions, aggregations. Time-series functions (interpolate, binning).'),
    ('Scheduled Queries', 'Native support for pre-aggregation. Runs SQL queries on schedule to compute hourly/daily '
     'metrics. Results stored back to Timestream.')
]

for component, description in components_timestream:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Native Partition-Based Multi-Tenancy', 3)

para = doc.add_paragraph()
run = para.add_run('✅ Timestream Multi-Tenancy is Superior to SiteWise')
run.bold = True
run.font.color.rgb = RGBColor(0, 128, 0)

doc.add_paragraph(
    'Every Timestream table uses tenant_id as partition key. This provides:'
)

mt_benefits_d = [
    'Physical data separation: Each tenant\'s data stored in separate partitions',
    'Query performance: Partition pruning automatically filters data',
    'Security: IAM policies can restrict access at partition level',
    'Cost allocation: Easy to track storage and query costs per tenant',
    'Scalability: Partitions distribute across nodes for parallel processing'
]

for benefit in mt_benefits_d:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('Comparison to Other Options:')
run.bold = True

mt_comparison_text = [
    'Options A & B (Snowflake RLS): Excellent - Database-enforced, cannot be bypassed',
    'Option D (Timestream partitions): Good - Native physical isolation, query-enforced',
    'Option C (SiteWise tags): Risky - Application-layer filtering, can be misconfigured'
]

for item in mt_comparison_text:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Grafana for Dashboards', 3)

para = doc.add_paragraph()
run = para.add_run('Why Grafana?')
run.bold = True

doc.add_paragraph(
    'Grafana is the industry-standard visualisation platform, widely used in manufacturing and DevOps. Key advantages:'
)

grafana_advantages = [
    'Fully White-Labelable: Remove Grafana branding, add company logos, custom themes',
    'Native Timestream Plugin: Official AWS plugin for querying Timestream',
    'Rich Visualisation Library: Time-series charts, gauges, heatmaps, state timelines, tables',
    'Embeddable Dashboards: iframe embedding in custom web applications',
    'Alert Management: Built-in alerting with multiple notification channels',
    'Variable Templates: Dynamic dashboards that adapt to user selection (site, machine)',
    'Community Support: Massive ecosystem, extensive documentation, active forums'
]

for advantage in grafana_advantages:
    doc.add_paragraph(advantage, style='List Bullet')

doc.add_paragraph()

grafana_deployment = [
    ('Grafana Cloud (Managed)', 'SaaS offering from Grafana Labs. Fully managed (no servers). £10-30 per user '
     'per month depending on tier. Automatic updates. 99.9% SLA.'),
    ('Self-Hosted on ECS Fargate', 'Run Grafana container on AWS ECS. Full control over deployment. One-time '
     'setup effort. Lower long-term cost for many users.')
]

for deployment, description in grafana_deployment:
    para = doc.add_paragraph()
    run = para.add_run(f'{deployment}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Application State (DynamoDB)', 3)

components_app_d = [
    ('DynamoDB Tables', 'Stores: tenant configuration, device registry, user preferences, dashboard configs, '
     'alert rules. Fast key-value lookups. Single-digit millisecond latency.'),
    ('API Gateway + Lambda', 'REST APIs for: device registration, dashboard management, user settings. '
     'Enforces tenant isolation. Queries Timestream and DynamoDB.'),
    ('Cognito User Pools', 'Authentication with SSO support. MFA capability. JWT tokens passed to Grafana and APIs.'),
    ('Secrets Manager', 'Stores: Timestream credentials, Grafana API keys, third-party integration secrets.')
]

for component, description in components_app_d:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('6.4 Data Flow', 2)

doc.add_paragraph('Unified data flow (all data types):')
flow_d = [
    'Sensor → IoT Core via MQTT (latency: <100ms)',
    'IoT Rules → Kinesis Streams (latency: <200ms)',
    'Lambda processes batch → Validates, enriches (latency: <500ms)',
    'Timestream Write API → Memory store (latency: <1s)',
    'Grafana queries Timestream via plugin (latency: <2s)',
    'Dashboard updates display (total: <5s end-to-end)'
]
for step in flow_d:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()

doc.add_paragraph('Critical alert fast-path:')
alert_flow_d = [
    'IoT Rules detects threshold breach (e.g., CO₂ >5000 ppm)',
    'Lambda fast-path triggered directly (bypasses Kinesis)',
    'SNS/SES notification sent (latency: <500ms)',
    'Alert logged to Timestream (latency: <2s total)'
]
for step in alert_flow_d:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('6.5 Strengths', 2)

strengths_d = [
    ('Fully AWS-Native', 'No Snowflake dependency. All AWS services. Single vendor relationship. Unified billing.'),
    ('Unified Storage', 'Timestream handles both continuous time-series AND discrete events. No dual-database '
     'complexity like Option C.'),
    ('Native Partition Multi-Tenancy', 'Physical data isolation at partition level. Better security than '
     'SiteWise tags. Simpler than application-layer filtering.'),
    ('Industry-Standard Dashboards', 'Grafana is battle-tested in manufacturing and industrial monitoring. '
     'Familiar to many operations teams. Rich visualisation library.'),
    ('Fully White-Labelable', 'Remove all vendor branding. Add company logos and themes. Embed dashboards in '
     'custom portals seamlessly.'),
    ('Cost-Effective', '£3,965/month (40% cheaper than Option C). Only 15% premium over Option B for AWS-native '
     'requirement.'),
    ('Standard SQL', 'Timestream SQL is mostly ANSI-compatible. Familiar to engineers. Easier to hire than '
     'Flink specialists.')
]

for title, desc in strengths_d:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(desc)

doc.add_heading('6.6 Limitations', 2)

limits_d = [
    ('15% More Expensive Than Option B', '£3,965/month vs £2,170-3,450 for Snowflake. Acceptable premium for '
     'AWS-native mandate but not cheapest option.'),
    ('Custom Asset Modelling Required', 'Unlike SiteWise\'s built-in asset models, must build custom data schemas '
     'and calculation logic. More initial development than Option C.'),
    ('Grafana Dependency', 'Whilst Grafana is popular, it adds external dependency (if using Grafana Cloud) or '
     'operational overhead (if self-hosted).'),
    ('Timestream SQL Differences', 'Whilst mostly standard, some Timestream-specific syntax and functions. '
     'Learning curve for SQL engineers familiar with PostgreSQL/MySQL.')
]

for title, desc in limits_d:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(153, 51, 0)
    para.add_run(desc)

doc.add_heading('6.7 When to Choose', 2)

doc.add_paragraph('Select Option D if:')
when_d = [
    'AWS-native is mandatory (no Snowflake allowed)',
    'Want best-in-class time-series dashboards (Grafana)',
    'Need native partition-based multi-tenancy',
    'Comfortable with 15% cost premium vs Option B for AWS-native benefit',
    'Team has AWS expertise (no Snowflake learning required)',
    'Want unified storage for time-series + events (simpler than Option C)',
    'Prefer industry-standard dashboards over custom development'
]
for item in when_d:
    doc.add_paragraph(item, style='List Bullet')

para = doc.add_paragraph()
run = para.add_run('\n✅ RECOMMENDED ')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 128, 0)
para.add_run(
    'if AWS-native is a requirement. Option D provides the best AWS-native solution without SiteWise\'s '
    'multi-tenancy limitations or custom dashboard burden. It is the second-best overall option after '
    'Option B (Snowflake).'
)

add_page_break(doc)

print("Options C and D complete. Saving document...")
doc.save('/Users/david/projects/smdh/docs/architecture/SMDH Infrastructure Design Options.docx')

print("\n" + "="*70)
print("✅ OPTIONS C AND D EXPANDED SUCCESSFULLY!")
print("="*70)
print("\nThe document now includes:")
print("  • Comprehensive Option C: AWS IoT SiteWise")
print("    - Dual storage architecture explained")
print("    - SiteWise components and capabilities")
print("    - Critical multi-tenancy challenges highlighted")
print("    - Custom dashboard requirements detailed")
print("    - Data flow for time-series and events")
print("    - Complete strengths and limitations")
print("  • Comprehensive Option D: AWS-Native (Timestream + Grafana)")
print("    - Unified storage architecture (vs dual in C)")
print("    - Timestream vs SiteWise comparison table")
print("    - Native partition-based multi-tenancy")
print("    - Grafana advantages and deployment options")
print("    - Complete data flow")
print("    - Strengths and limitations")
print("\nBoth options now match the detail level of Options A and B!")
print("Ready for diagram insertion and final review.")

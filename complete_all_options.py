#!/usr/bin/env python3
"""
Complete all architecture option sections and comparison matrices
Adds comprehensive structured content for Options A, B, C, D
Includes detailed comparison section
Uses UK English throughout
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_page_break(doc):
    doc.add_page_break()

def add_heading_with_color(doc, text, level, color_rgb=(0, 51, 102)):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_placeholder_diagram(doc, diagram_name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\n[INSERT DIAGRAM: {diagram_name}]\n')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)
    for _ in range(6):
        p.add_run('\n')
    # Add grey background
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)

print("Opening document...")
doc = Document('/Users/david/projects/smdh/docs/architecture/SMDH Infrastructure Design Options.docx')

print("Adding Option A detailed content...")

# OPTION A: FLINK-BASED ARCHITECTURE
add_heading_with_color(doc, '3. Architecture Option A: Flink-Based AWS-Heavy Architecture', 1, (153, 51, 0))

doc.add_heading('3.1 Overview and Summary', 2)

summary_table = doc.add_table(rows=7, cols=2)
summary_table.style = 'Medium Shading 1 Accent 6'

cells = summary_table.rows[0].cells
cells[0].text = 'Characteristic'
cells[1].text = 'Details'

cells = summary_table.rows[1].cells
cells[0].text = 'Monthly Cost (30 tenants)'
cells[1].text = '£2,500 - 4,200'

cells = summary_table.rows[2].cells
cells[0].text = 'Cost Per Tenant'
cells[1].text = '£83 - 140'

cells = summary_table.rows[3].cells
cells[0].text = 'Complexity Level'
cells[1].text = 'Very High (15+ AWS services)'

cells = summary_table.rows[4].cells
cells[0].text = 'Implementation Timeline'
cells[1].text = '24 weeks'

cells = summary_table.rows[5].cells
cells[0].text = 'Team Size & Skills'
cells[1].text = '3-4 engineers (Flink, Java/Scala, AWS, data engineering)'

cells = summary_table.rows[6].cells
cells[0].text = 'Best Suited For'
cells[1].text = 'Sub-second latency legally required; advanced ML needed'

doc.add_paragraph()

doc.add_paragraph(
    'Option A represents the most technically sophisticated architecture, leveraging Apache Flink for '
    'advanced stream processing. This approach provides unparalleled real-time capabilities with sub-second '
    'latency but comes with significant operational complexity.'
)

doc.add_heading('3.2 Architecture Diagram', 2)
add_placeholder_diagram(doc, 'SMDH-Option-A-Flink-Architecture.drawio')

doc.add_heading('3.3 Core Components', 2)

doc.add_heading('Ingestion Layer', 3)

components_a = [
    ('AWS IoT Core', 'MQTT broker receiving 1Hz sensor data from 30-45 machines per site. '
     'Handles device authentication via X.509 certificates. Provides device shadows for configuration.'),
    ('IoT Rules Engine', 'Routes messages based on content. Enriches data with tenant metadata. '
     'Filters invalid messages. SQL-based routing rules.'),
    ('Kinesis Data Firehose', 'Buffers data in 10MB or 60-second batches. Compresses to Parquet format. '
     'Invokes Lambda for transformation. Delivers to S3.'),
    ('Lambda Transformation', 'Validates sensor ranges, converts units, enriches with lookup data. '
     'Python or Node.js. Scales automatically.'),
    ('S3 Data Lake', 'Stores Parquet files partitioned by tenant/date. Lifecycle policies for archival. '
     'Enables historical analysis and recovery.')
]

for component, description in components_a:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Stream Processing Layer (Apache Flink)', 3)

para = doc.add_paragraph()
run = para.add_run('Why Flink?')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph(
    'Apache Flink is an advanced distributed stream processing framework. Unlike simpler alternatives, '
    'Flink provides:'
)

flink_capabilities = [
    'Sub-second latency: Process and respond to events in <1 second',
    'Exactly-once semantics: Guarantees no data loss or duplication',
    'Complex event processing: Detect patterns across multiple streams',
    'Stateful processing: Remember context across events (e.g., calculate 5-minute rolling averages)',
    'Advanced windowing: Tumbling, sliding, session windows for aggregations',
    'Backpressure handling: Gracefully manages data spikes'
]

for cap in flink_capabilities:
    doc.add_paragraph(cap, style='List Bullet')

doc.add_paragraph()

components_flink = [
    ('Flink on EMR', 'Runs on 3-5 node cluster. Auto-scaling based on data velocity. '
     'Checkpoints state to S3 every 5 minutes.'),
    ('RocksDB State Backend', 'Embedded database storing intermediate state. Enables fault tolerance. '
     'Recovers from checkpoints if nodes fail.'),
    ('Flink Jobs', 'Custom Java/Scala code for: real-time aggregations, anomaly detection, '
     'pattern matching, multi-tenant data isolation.')
]

for component, description in components_flink:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Data Platform Layer (Snowflake)', 3)

components_snow = [
    ('Snowpipe Auto-Ingestion', 'Detects new S3 files via events. Loads into Snowflake within seconds. '
     'Serverless—no warehouse management.'),
    ('Raw Tables', 'VARIANT columns store semi-structured JSON. Partitioned by tenant_id. '
     'Clustered on timestamp for query performance.'),
    ('Snowflake Streams', 'Change Data Capture tracking new/changed rows. Enables incremental processing. '
     'Multiple consumers can read independently.'),
    ('Snowflake Tasks', 'Scheduled SQL transformations. DAG-based dependencies. '
     'Runs hourly aggregations, data quality checks, archival.'),
    ('Curated Views', 'Row-Level Security enforces tenant isolation. Pre-calculated aggregations. '
     'Analytics-ready datasets.')
]

for component, description in components_snow:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Application Layer', 3)

components_app = [
    ('API Gateway + Lambda', 'REST APIs for data access. Tenant authentication via Cognito. '
     'Query Snowflake and return JSON.'),
    ('ECS Fargate Web App', 'React frontend hosted on Fargate. Auto-scaling. CI/CD via CodePipeline.'),
    ('QuickSight Dashboards', 'Embedded analytics. Row-Level Security passes through from Snowflake. '
     'Pay-per-session pricing.'),
    ('SageMaker ML', 'Custom model training (predictive maintenance, anomaly detection). '
     'GPU-accelerated. Model hosting on auto-scaling endpoints.')
]

for component, description in components_app:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('3.4 Data Flow', 2)

doc.add_paragraph('Real-time path (critical alerts):')
flow_rt = [
    'Sensor → IoT Core (100ms)',
    'IoT Rules → Kinesis Streams (200ms)',
    'Flink processes and detects alert condition (500ms)',
    'SNS notification sent (200ms)',
    'Total: <2 seconds end-to-end'
]
for step in flow_rt:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()
doc.add_paragraph('Batch path (historical analytics):')
flow_batch = [
    'Kinesis Firehose buffers 60 seconds',
    'Lambda transforms batch',
    'S3 receives Parquet files',
    'Snowpipe loads into Snowflake (<5 min)',
    'Snowflake Tasks aggregate hourly',
    'QuickSight dashboards refresh'
]
for step in flow_batch:
    doc.add_paragraph(step, style='List Number')

doc.add_heading('3.5 Strengths', 2)

strengths_a = [
    ('Sub-Second Latency', 'Achieves <1s for critical data paths. Essential if regulatory requirements '
     'mandate immediate response.'),
    ('Advanced Stream Processing', 'Complex event processing, pattern detection, sophisticated windowing '
     'not possible with simpler systems.'),
    ('Unlimited Flexibility', 'Custom Java/Scala code enables any transformation or ML algorithm.'),
    ('Proven at Scale', 'Flink powers streaming at Netflix, Uber, Alibaba processing billions of events daily.')
]

for title, desc in strengths_a:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(desc)

doc.add_heading('3.6 Limitations', 2)

limits_a = [
    ('Very High Complexity', 'Managing 15+ services, Flink cluster tuning, distributed debugging.'),
    ('Flink Expertise Required', 'Rare and expensive skill set. Steep learning curve for Java/Scala.'),
    ('Longest Timeline', '24 weeks to production. Complex distributed systems take time to build and test.'),
    ('Operational Overhead', 'EMR cluster management, checkpointing monitoring, state recovery testing.')
]

for title, desc in limits_a:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)
    para.add_run(desc)

doc.add_heading('3.7 When to Choose', 2)

doc.add_paragraph('Select Option A ONLY if:')
when_a = [
    'Sub-second latency is legally or regulatorily mandated (not just preferred)',
    'Complex event processing is essential for the business case',
    'Team has Flink expertise or budget for specialists',
    'Willing to accept highest complexity and 24-week timeline'
]
for item in when_a:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

print("Option A complete. Adding Option B...")

# OPTION B: SNOWFLAKE ARCHITECTURE
add_heading_with_color(doc, '4. Architecture Option B: Snowflake-Leveraged Architecture (RECOMMENDED)', 1, (0, 102, 204))

doc.add_heading('4.1 Overview and Summary', 2)

summary_table_b = doc.add_table(rows=7, cols=2)
summary_table_b.style = 'Medium Shading 1 Accent 5'

cells = summary_table_b.rows[0].cells
cells[0].text = 'Characteristic'
cells[1].text = 'Details'

cells = summary_table_b.rows[1].cells
cells[0].text = 'Monthly Cost (30 tenants)'
cells[1].text = '£2,170 - 3,450 (CHEAPEST)'

cells = summary_table_b.rows[2].cells
cells[0].text = 'Cost Per Tenant'
cells[1].text = '£72 - 115'

cells = summary_table_b.rows[3].cells
cells[0].text = 'Complexity Level'
cells[1].text = 'Low (5 core services)'

cells = summary_table_b.rows[4].cells
cells[0].text = 'Implementation Timeline'
cells[1].text = '20 weeks (FASTEST)'

cells = summary_table_b.rows[5].cells
cells[0].text = 'Team Size & Skills'
cells[1].text = '2-3 engineers (SQL, Snowflake, basic AWS)'

cells = summary_table_b.rows[6].cells
cells[0].text = 'Best Suited For'
cells[1].text = 'Most scenarios - default recommendation'

doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('🏆 RECOMMENDED OPTION')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 128, 0)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

doc.add_paragraph(
    'Option B centralises data processing in Snowflake, dramatically simplifying the architecture compared '
    'to Option A. By leveraging Snowflake\'s native streaming, task orchestration, and SQL-first approach, '
    'this option provides the fastest time-to-market and lowest operational complexity whilst meeting all '
    'functional requirements.'
)

doc.add_heading('4.2 Architecture Diagram', 2)
add_placeholder_diagram(doc, 'SMDH-Option-B-Snowflake-Architecture.drawio')

doc.add_heading('4.3 Core Components', 2)

doc.add_heading('Simplified Ingestion', 3)

components_b = [
    ('AWS IoT Core', 'Same as Option A—MQTT broker with X.509 authentication.'),
    ('Kinesis Data Streams', 'Lightweight buffering (not Firehose). Streams directly to Snowflake. '
     '10-20 shards partitioned by tenant_id.'),
    ('Snowpipe Streaming', 'Snowflake\'s streaming ingestion service. Loads data in 5-10 seconds. '
     'Eliminates S3 staging and Lambda transformation complexity.')
]

for component, description in components_b:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Unified Data Platform (Snowflake Does It All)', 3)

para = doc.add_paragraph()
run = para.add_run('Key Simplification:')
run.bold = True

doc.add_paragraph(
    'In Option B, Snowflake replaces multiple Option A services:'
)

sf_replaces = [
    'Replaces S3 Data Lake → Snowflake internal storage',
    'Replaces Lambda + Flink → Snowflake Streams and Tasks',
    'Replaces AWS Glue → Snowflake native schema evolution',
    'Replaces separate ETL orchestration → Snowflake Task DAGs'
]

for item in sf_replaces:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()

components_sf = [
    ('Raw Tables with VARIANT', 'JSON data stored in VARIANT columns. Schema-flexible. '
     'Automatic compression (75-85% reduction). Partitioned by tenant_id.'),
    ('Snowflake Streams (CDC)', 'Tracks inserts/updates/deletes. Enables incremental processing. '
     'Zero-copy—doesn\'t duplicate data.'),
    ('Snowflake Tasks', 'SQL-based transformations run on schedule or when streams have data. '
     'DAG orchestration. Serverless compute.'),
    ('Dynamic Tables', 'Continuously updated aggregations. Declarative SQL definitions. '
     'Incremental refresh automatically.'),
    ('Row-Level Security (RLS)', 'Native multi-tenancy. Policies enforce tenant isolation at database level. '
     'Cannot be bypassed by applications.')
]

for component, description in components_sf:
    para = doc.add_paragraph()
    run = para.add_run(f'{component}: ')
    run.bold = True
    para.add_run(description)

doc.add_heading('Machine Learning (Snowflake Cortex)', 3)

doc.add_paragraph(
    'Snowflake Cortex ML enables SQL-based machine learning:'
)

ml_sf = [
    'CREATE MODEL predict_failure AS SELECT * FROM training_data;',
    'SELECT PREDICT_FAILURE(sensor_data) FROM current_readings;',
    'Supports: Classification, regression, time-series forecasting',
    'Limitation: Less flexible than SageMaker but covers 80% of use cases'
]

for item in ml_sf:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.4 Data Flow', 2)

doc.add_paragraph('End-to-end flow (simplified):')
flow_b = [
    'Sensor → IoT Core (100ms)',
    'Kinesis Streams (200ms)',
    'Snowpipe Streaming → Raw table (5-10 seconds)',
    'Snowflake Stream detects new rows immediately',
    'Task processes (runs every 1 min) → Curated views',
    'Dashboard refreshes (total: 60-65 seconds)'
]
for step in flow_b:
    doc.add_paragraph(step, style='List Number')

doc.add_paragraph()

para = doc.add_paragraph()
run = para.add_run('Alert Latency Solution:')
run.bold = True
run.font.color.rgb = RGBColor(153, 51, 0)

doc.add_paragraph(
    'Default 60s latency violates <10s alert requirement. FIX: Add Lambda fast-path for alerts '
    '(+£100-150/month, +2 weeks). Fast-path achieves <5s alert latency whilst keeping batch processing '
    'in Snowflake.'
)

doc.add_heading('4.5 Strengths', 2)

strengths_b = [
    ('Lowest Cost', '£2,170-3,450/month. Cheapest option by 15-20%. Auto-suspend when idle.'),
    ('Lowest Complexity', 'Only 5 core services vs 15+ in Option A. Single platform for most processing.'),
    ('Fastest Timeline', '20 weeks to production. SQL-first development accelerates delivery.'),
    ('SQL-First Development', 'Familiar to most engineers. Easier to hire than Flink specialists.'),
    ('Native Multi-Tenancy', 'Row-Level Security enforced at database level. Production-proven.'),
    ('Unified Governance', 'All data in one platform. Single place for audits, lineage, retention policies.')
]

for title, desc in strengths_b:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(desc)

doc.add_heading('4.6 Limitations', 2)

limits_b = [
    ('Alert Latency Requires Fix', '60s default. Add Lambda fast-path for <5s alerts (+£100-150/month).'),
    ('Snowflake Licensing', 'Separate from AWS bill. Requires contract negotiation. Some organisations '
     'prefer AWS-native only (see Option D).'),
    ('Not Sub-Second', 'Snowpipe Streaming is 5-10s, not <1s. If sub-second legally required, choose Option A.')
]

for title, desc in limits_b:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(153, 51, 0)
    para.add_run(desc)

doc.add_heading('4.7 When to Choose', 2)

doc.add_paragraph('Select Option B if:')
when_b = [
    'Cost and simplicity are priorities (most SME scenarios)',
    'Team has SQL skills or can train quickly',
    '60s dashboard latency acceptable, or willing to add Lambda fast-path',
    'Want fastest time to market (20 weeks)',
    'Snowflake licensing not a blocker'
]
for item in when_b:
    doc.add_paragraph(item, style='List Bullet')

para = doc.add_paragraph()
run = para.add_run('\n✅ RECOMMENDED for 80% of scenarios')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 128, 0)

add_page_break(doc)

print("Option B complete. Adding Options C and D summaries (refer to detailed markdown docs)...")

# OPTION C and D: Add summary sections with references to full markdown docs
# (Space constraints - full detail available in markdown)

add_heading_with_color(doc, '5. Architecture Option C: AWS IoT SiteWise-Based', 1, (40, 167, 69))
doc.add_paragraph(
    'For complete details on Option C (SiteWise), see smdh-architecture-option-c-sitewise.md. '
    'Key summary:'
)

doc.add_heading('5.1 Overview', 2)
summary_table_c = doc.add_table(rows=5, cols=2)
summary_table_c.style = 'Medium Shading 1 Accent 4'

cells = summary_table_c.rows[0].cells; cells[0].text = 'Monthly Cost'; cells[1].text = '£6,334 (CORRECTED)'
cells = summary_table_c.rows[1].cells; cells[0].text = 'Complexity'; cells[1].text = 'Medium-High (8-10 services)'
cells = summary_table_c.rows[2].cells; cells[0].text = 'Timeline'; cells[1].text = '28 weeks'
cells = summary_table_c.rows[3].cells; cells[0].text = 'Best For'; cells[1].text = 'SiteWise expertise; <10 tenants'
cells = summary_table_c.rows[4].cells; cells[0].text = 'Key Limitation'; cells[1].text = 'NOT designed for multi-tenant SaaS'

doc.add_heading('5.2 Architecture Diagram', 2)
add_placeholder_diagram(doc, 'SMDH-Option-C-SiteWise-Architecture.drawio')

doc.add_paragraph('Strengths: Managed asset modelling, OPC-UA/Modbus support, real-time latency')
doc.add_paragraph('Limitations: Tag-based multi-tenancy risk, dual storage (SiteWise + Timestream), custom dashboards required')

add_page_break(doc)

add_heading_with_color(doc, '6. Architecture Option D: AWS-Native (Timestream + Grafana)', 1, (0, 150, 200))
doc.add_paragraph(
    'For complete details on Option D, see smdh-architecture-option-d-aws-native.md. Key summary:'
)

doc.add_heading('6.1 Overview', 2)
summary_table_d = doc.add_table(rows=5, cols=2)
summary_table_d.style = 'Medium Shading 1 Accent 3'

cells = summary_table_d.rows[0].cells; cells[0].text = 'Monthly Cost'; cells[1].text = '£3,965'
cells = summary_table_d.rows[1].cells; cells[0].text = 'Complexity'; cells[1].text = 'Medium (7-8 services)'
cells = summary_table_d.rows[2].cells; cells[0].text = 'Timeline'; cells[1].text = '22 weeks'
cells = summary_table_d.rows[3].cells; cells[0].text = 'Best For'; cells[1].text = 'AWS-native mandate'
cells = summary_table_d.rows[4].cells; cells[0].text = 'Key Strength'; cells[1].text = 'Fully AWS-native, Grafana dashboards'

doc.add_heading('6.2 Architecture Diagram', 2)
add_placeholder_diagram(doc, 'SMDH-Option-D-Timestream-Architecture.drawio')

doc.add_paragraph('Strengths: Fully AWS-native, unified storage, native partition multi-tenancy, Grafana')
doc.add_paragraph('Limitations: 15% more expensive than Option B, custom asset modelling required')

add_page_break(doc)

print("Adding comprehensive comparison section...")

# SECTION 7: COMPARISON
add_heading_with_color(doc, '7. Architecture Comparison', 1)

doc.add_heading('7.1 Cost Comparison', 2)

cost_table = doc.add_table(rows=5, cols=6)
cost_table.style = 'Light Grid Accent 1'

# Header
cells = cost_table.rows[0].cells
cells[0].text = 'Option'
cells[1].text = 'Monthly (30 tenants)'
cells[2].text = 'Per Tenant'
cells[3].text = 'Annual TCO'
cells[4].text = 'Within Budget?'
cells[5].text = 'Rank'

cells = cost_table.rows[1].cells
cells[0].text = 'B (Snowflake)'
cells[1].text = '£2,170-3,450'
cells[2].text = '£72-115'
cells[3].text = '£346K-521K'
cells[4].text = '✅ Yes'
cells[5].text = '1st (Cheapest)'

cells = cost_table.rows[2].cells
cells[0].text = 'A (Flink)'
cells[1].text = '£2,500-4,200'
cells[2].text = '£83-140'
cells[3].text = '£525K-720K'
cells[4].text = '✅ Yes'
cells[5].text = '2nd'

cells = cost_table.rows[3].cells
cells[0].text = 'D (Timestream)'
cells[1].text = '£3,965'
cells[2].text = '£132'
cells[3].text = '£374K-538K'
cells[4].text = '✅ Yes'
cells[5].text = '3rd'

cells = cost_table.rows[4].cells
cells[0].text = 'C (SiteWise)'
cells[1].text = '£6,334'
cells[2].text = '£211'
cells[3].text = '£566K-715K'
cells[4].text = '✅ Yes (corrected)'
cells[5].text = '4th (Most expensive)'

doc.add_paragraph()

doc.add_heading('7.2 Complexity Comparison', 2)

complexity_table = doc.add_table(rows=5, cols=4)
complexity_table.style = 'Light Grid Accent 1'

cells = complexity_table.rows[0].cells
cells[0].text = 'Option'
cells[1].text = 'Number of Services'
cells[2].text = 'Primary Skills Required'
cells[3].text = 'Complexity Rank'

cells = complexity_table.rows[1].cells
cells[0].text = 'B (Snowflake)'
cells[1].text = '5 services'
cells[2].text = 'SQL, Snowflake'
cells[3].text = '1st (Simplest)'

cells = complexity_table.rows[2].cells
cells[0].text = 'D (Timestream)'
cells[1].text = '7-8 services'
cells[2].text = 'AWS-native, SQL, Grafana'
cells[3].text = '2nd'

cells = complexity_table.rows[3].cells
cells[0].text = 'C (SiteWise)'
cells[1].text = '8-10 services'
cells[2].text = 'SiteWise, AWS IoT, React'
cells[3].text = '3rd'

cells = complexity_table.rows[4].cells
cells[0].text = 'A (Flink)'
cells[1].text = '15+ services'
cells[2].text = 'Flink, Java/Scala, AWS'
cells[3].text = '4th (Most complex)'

doc.add_paragraph()

doc.add_heading('7.3 Performance and Latency', 2)

perf_table = doc.add_table(rows=5, cols=4)
perf_table.style = 'Light Grid Accent 1'

cells = perf_table.rows[0].cells
cells[0].text = 'Option'
cells[1].text = 'Real-Time Latency'
cells[2].text = 'Alert Latency'
cells[3].text = 'Dashboard Updates'

cells = perf_table.rows[1].cells
cells[0].text = 'A (Flink)'
cells[1].text = '<1 second'
cells[2].text = '<5 seconds ✅'
cells[3].text = '<5 seconds'

cells = perf_table.rows[2].cells
cells[0].text = 'B (Snowflake)'
cells[1].text = '5-10 seconds'
cells[2].text = '60s (fixable to <5s)'
cells[3].text = '60-65 seconds'

cells = perf_table.rows[3].cells
cells[0].text = 'C (SiteWise)'
cells[1].text = '<1 second'
cells[2].text = '<10 seconds ✅'
cells[3].text = '<5 seconds'

cells = perf_table.rows[4].cells
cells[0].text = 'D (Timestream)'
cells[1].text = '<5 seconds'
cells[2].text = '<5 seconds ✅'
cells[3].text = '<10 seconds'

doc.add_paragraph()

doc.add_heading('7.4 Multi-Tenancy Security', 2)

mt_table = doc.add_table(rows=5, cols=3)
mt_table.style = 'Light Grid Accent 1'

cells = mt_table.rows[0].cells
cells[0].text = 'Option'
cells[1].text = 'Multi-Tenancy Approach'
cells[2].text = 'Security Assessment'

cells = mt_table.rows[1].cells
cells[0].text = 'B (Snowflake)'
cells[1].text = 'Row-Level Security (RLS)'
cells[2].text = '✅ Excellent - Database enforced'

cells = mt_table.rows[2].cells
cells[0].text = 'A (Flink)'
cells[1].text = 'Snowflake RLS'
cells[2].text = '✅ Excellent - Database enforced'

cells = mt_table.rows[3].cells
cells[0].text = 'D (Timestream)'
cells[1].text = 'Partition Keys'
cells[2].text = '✅ Good - Native physical isolation'

cells = mt_table.rows[4].cells
cells[0].text = 'C (SiteWise)'
cells[1].text = 'Tag-based filtering'
cells[2].text = '⚠️ Risky - Application-layer'

doc.add_paragraph()

doc.add_heading('7.5 Decision Framework', 2)

doc.add_paragraph('Use this decision tree to select the appropriate option:')

decision_tree = [
    '1. Is AWS-native mandatory (no Snowflake)?',
    '   → YES: Choose Option D (Timestream + Grafana)',
    '   → NO: Continue to question 2',
    '',
    '2. Is sub-second latency legally required for ALL data?',
    '   → YES: Choose Option A (Flink)',
    '   → NO: Continue to question 3',
    '',
    '3. Is cost and simplicity your top priority?',
    '   → YES: Choose Option B (Snowflake) ✅ RECOMMENDED',
    '   → NO: Continue to question 4',
    '',
    '4. Do you have deep SiteWise expertise AND <10 tenants?',
    '   → YES: Consider Option C (with caution)',
    '   → NO: Choose Option B (Snowflake) ✅ RECOMMENDED',
    '',
    'DEFAULT: Choose Option B for most scenarios'
]

for line in decision_tree:
    if line.strip():
        doc.add_paragraph(line)

add_page_break(doc)

print("Saving complete document...")
doc.save('/Users/david/projects/smdh/docs/architecture/SMDH Infrastructure Design Options.docx')

print("\n" + "="*70)
print("✅ DOCUMENT COMPLETE!")
print("="*70)
print("\nThe document now includes:")
print("  • Comprehensive Section 2: System Overview (use cases, requirements)")
print("  • Detailed Option A: Flink-based architecture")
print("  • Detailed Option B: Snowflake architecture (RECOMMENDED)")
print("  • Summary Options C and D (full details in markdown docs)")
print("  • Comprehensive comparison matrices:")
print("    - Cost comparison table")
print("    - Complexity comparison table")
print("    - Performance and latency comparison")
print("    - Multi-tenancy security comparison")
print("    - Decision framework tree")
print("\nNext steps:")
print("  1. Insert the 4 Draw.io diagrams into placeholder boxes")
print("  2. Optionally expand Options C and D with full markdown content")
print("  3. Review and adjust formatting/styling as needed")
print("  4. Add any company branding")

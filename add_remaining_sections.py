#!/usr/bin/env python3
"""
Continue building the comprehensive SMDH Architecture document
Adds Options C, D, Comparison section, and Glossary
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def add_heading_with_color(doc, text, level, color_rgb=(0, 51, 102)):
    """Add a coloured heading"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_placeholder_box(doc, text, height=3.0):
    """Add a placeholder box for diagrams"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add shaded box
    run = p.add_run(f'\n{text}\n')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Add some spacing
    for _ in range(int(height * 2)):
        p.add_run('\n')

    # Add border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '24')
        border_el.set(qn('w:space'), '4')
        border_el.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border_el)
    pPr.append(pBdr)

    # Add shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)

    return p

# Open existing document
doc = Document('/Users/david/projects/smdh/docs/architecture/SMDH-Infrastructure-Design-Options-v2.docx')

print("Loaded existing document. Adding Option C...")

#=============================================================================
# SECTION 6: OPTION C - AWS IOT SITEWISE ARCHITECTURE
#=============================================================================

add_heading_with_color(doc, '6. Architecture Option C: AWS IoT SiteWise-Based Architecture', 1, (40, 167, 69))

doc.add_heading('6.1 Overview', 2)

# Summary box
summary_table = doc.add_table(rows=6, cols=2)
summary_table.style = 'Medium Shading 1 Accent 4'

cells = summary_table.rows[0].cells
cells[0].text = 'Monthly Cost (30 tenants)'
cells[1].text = '£6,334 ✅ CORRECTED (was £21,150)'

cells = summary_table.rows[1].cells
cells[0].text = 'Cost Per Tenant'
cells[1].text = '£211'

cells = summary_table.rows[2].cells
cells[0].text = 'Complexity Level'
cells[1].text = 'Medium-High (8-10 services)'

cells = summary_table.rows[3].cells
cells[0].text = 'Implementation Timeline'
cells[1].text = '28 weeks (longest)'

cells = summary_table.rows[4].cells
cells[0].text = 'Team Size Required'
cells[1].text = '3-4 engineers (SiteWise + AWS IoT)'

cells = summary_table.rows[5].cells
cells[0].text = 'Best Suited For'
cells[1].text = 'Teams with deep SiteWise expertise; <10 tenants'

doc.add_paragraph()

doc.add_heading('What is This Architecture?', 3)

doc.add_paragraph(
    'Option C leverages AWS IoT SiteWise, a managed service specifically designed for industrial equipment monitoring. '
    'Imagine a platform purpose-built for factories—it "speaks" the language of industrial sensors and equipment, '
    'with built-in concepts like asset hierarchies (e.g., "Factory → Production Line → Machine → Sensor"), time-series '
    'storage, and automatic metric calculations.'
)

doc.add_paragraph(
    'Unlike general-purpose databases, SiteWise understands manufacturing concepts natively. You define an "asset model" '
    '(like a template for a CNC machine), and SiteWise automatically organises data, calculates metrics like Overall '
    'Equipment Effectiveness (OEE), and provides specialised dashboards.'
)

para = doc.add_paragraph()
run = para.add_run('\n⚠️ IMPORTANT COST CORRECTION\n')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(153, 51, 0)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'A previous analysis incorrectly calculated Option C cost at £21,150/month due to a 7.3x data volume error. '
    'The CORRECTED cost is £6,334/month (70% reduction), making this option economically viable. However, it remains '
    'the most expensive option and has multi-tenancy limitations.'
)

doc.add_heading('6.2 Architecture Diagram', 2)

add_placeholder_box(doc, '[INSERT: SMDH-Option-C-SiteWise-Architecture.drawio diagram here]', 4.0)

doc.add_paragraph(
    'The diagram shows a dual-storage approach: SiteWise for continuous sensor data and Timestream for discrete '
    'events, with a custom React application providing white-labelled dashboards.'
)

doc.add_heading('6.3 Detailed Component Description', 2)

doc.add_heading('Layer 1: Dual Ingestion Path', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('AWS IoT SiteWise Gateway (On-Premises)\n')
run.bold = True

doc.add_paragraph(
    'What it does: A software appliance (runs on a local server at the factory) that connects to industrial protocols '
    'like OPC-UA (common in manufacturing) and Modbus (used by energy metres). It "translates" these protocols into '
    'AWS IoT format.\n\n'
    'Why it matters: Many manufacturing sensors don\'t speak MQTT directly. The gateway bridges legacy equipment to '
    'the cloud without replacing existing infrastructure.\n\n'
    'Technical details: Runs as Docker containers on EC2 instances or on-premises hardware. Supports OPC-UA, Modbus TCP, '
    'Ethernet/IP. Includes local caching for offline operation and automatic retry.\n\n'
    'Operational consideration: Requires managing on-premises software at each manufacturing site. Updates, security '
    'patches, and troubleshooting must be performed remotely or on-site.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Split Routing: Time-Series vs Events\n')
run.bold = True

doc.add_paragraph(
    'What it does: IoT Rules Engine routes data to different destinations based on type:\n'
    '  • Continuous data (machine status, temperature) → SiteWise\n'
    '  • Discrete events (RFID scans, job completions) → Kinesis → Timestream\n\n'
    'Why it matters: SiteWise is optimised for time-series data, not discrete events. Job tracking requires a different '
    'storage model (Timestream), creating a dual-storage architecture.\n\n'
    'Complexity impact: Applications must query two different databases. Data correlation across systems requires custom '
    'logic. Increases overall system complexity.'
)

doc.add_heading('Layer 2: AWS IoT SiteWise (Core Platform)', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('SiteWise Asset Models and Hierarchies\n')
run.bold = True
run.font.color.rgb = RGBColor(40, 167, 69)

doc.add_paragraph(
    'What it does: Define templates for equipment types (e.g., "CNC Machine Model X") with properties (spindle speed, '
    'temperature), metrics (utilisation %), and alarms. Organise assets into hierarchies (Site → Building → Floor → '
    'Machine).\n\n'
    'Why it matters: Provides structure for industrial data. Calculations like OEE (Overall Equipment Effectiveness) '
    'are built-in formulas, not custom code. Hierarchy enables roll-up reporting (see total factory OEE, drill down '
    'to specific machines).\n\n'
    'Technical details: Asset models defined via API or console. Supports inheritance (base model → variants). Formulas '
    'use SiteWise Expression Language (proprietary, similar to Excel formulas). Assets can have static attributes '
    '(serial number, install date) and dynamic measurements.\n\n'
    'Manufacturing-specific benefit: Unlike generic databases, SiteWise "understands" concepts like availability, '
    'performance, and quality (the three components of OEE).'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('SiteWise Compute Expressions\n')
run.bold = True

doc.add_paragraph(
    'What it does: Automatically calculates derived metrics from raw sensor data. Example: "If machine_state = RUNNING, '
    'increment running_time; calculate utilisation = running_time / total_time."\n\n'
    'Why it matters: No Lambda functions needed for standard manufacturing calculations. SiteWise performs these '
    'computations in real-time as data arrives.\n\n'
    'Technical details: Transform metrics (unit conversions, averaging), Metric metrics (aggregations like sum/avg/min/max '
    'over windows), Measurement metrics (raw sensor values). Computed at ingestion time. Results stored alongside raw data.\n\n'
    'Limitation: Expression language less flexible than Python/Java. Complex business logic still requires Lambda.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('SiteWise Time-Series Storage\n')
run.bold = True

doc.add_paragraph(
    'What it does: Stores sensor readings in a managed time-series database optimised for industrial IoT query patterns.\n\n'
    'Why it matters: Automatic tiering: hot tier (13 months) for fast queries, cold tier (14+ months) for archival. '
    'No manual lifecycle management.\n\n'
    'Technical details: Storage charged per GB stored + per million values. Automatically indexes by timestamp and asset '
    'hierarchy for efficient queries. Supports batch and real-time queries.\n\n'
    'Cost structure: $0.75 per million values ingested + $0.35/GB/month hot tier + $0.03/GB/month cold tier. This pricing '
    'model can be expensive for high-frequency sensors (1Hz = 86,400 values/day/sensor).'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('⚠️ Multi-Tenancy Challenge: Tag-Based Isolation\n')
run.bold = True
run.font.color.rgb = RGBColor(153, 51, 0)

doc.add_paragraph(
    'What it does: SiteWise does NOT have native multi-tenancy. The workaround is to tag every asset with a "tenant_id" '
    'and filter queries in the application layer.\n\n'
    'Why this is problematic: Unlike Snowflake\'s Row-Level Security (database-enforced), tag-based filtering relies '
    'on application code correctly filtering queries. A bug or misconfiguration could expose one tenant\'s data to another.\n\n'
    'Technical implementation: Lambda functions must add WHERE tenant_id = :authorized_tenant to EVERY query. Custom '
    'React application must enforce filtering in UI. Requires extensive security testing.\n\n'
    'Risk assessment: Medium-high risk for SaaS platforms. Acceptable for single-tenant deployments or <10 trusted tenants. '
    'Requires rigorous code review and penetration testing for multi-tenant SaaS.'
)

doc.add_heading('Layer 3: Event Storage (Amazon Timestream)', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Why Dual Storage? SiteWise Limitations for Events\n')
run.bold = True

doc.add_paragraph(
    'SiteWise is designed for continuous sensor readings (machine running at 1500 RPM, temperature is 65°C). It is '
    'NOT well-suited for discrete events like:\n'
    '  • RFID scan at 10:23 AM: "Job #1234 scanned at Station 5"\n'
    '  • Barcode read: "Product X456 completed assembly"\n\n'
    'These events have irregular timing, rich metadata (job details, operator ID), and different query patterns (find '
    'all scans for Job #1234 in the past week). Timestream is purpose-built for this use case.\n\n'
    'Consequence: Option C requires BOTH SiteWise (time-series) AND Timestream (events), increasing complexity.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Amazon Timestream for Event Data\n')
run.bold = True

doc.add_paragraph(
    'What it does: Stores job tracking events, RFID scans, barcode reads, and other discrete event data.\n\n'
    'Why it matters: Native SQL interface. Efficient for event-based queries ("find all movements of Job #1234"). '
    'Separate from SiteWise to avoid forcing event data into time-series model.\n\n'
    'Technical details: Same as Option D (see Section 7). Memory store for recent data, magnetic store for historical. '
    'SQL query engine. Native multi-tenancy via partition keys (better than SiteWise\'s tags).'
)

doc.add_heading('Layer 4: Custom Dashboard Layer', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Why Custom Dashboards? SiteWise Monitor Limitations\n')
run.bold = True

doc.add_paragraph(
    'AWS provides SiteWise Monitor—a built-in dashboarding tool. However, it has critical limitations for SaaS:\n'
)

monitor_limits = [
    'Cannot fully white-label: AWS branding and URLs remain visible',
    'Limited customisation: Dashboard layout and styling options are constrained',
    'No multi-tenancy support: Not designed for SaaS platform with many companies',
    'Basic visualisations: Chart types limited compared to QuickSight or Grafana'
]

for item in monitor_limits:
    doc.add_paragraph(f'  • {item}')

doc.add_paragraph(
    '\nConsequence: Option C requires building a custom React application that calls SiteWise and Timestream APIs directly. '
    'This adds 6-8 weeks to the timeline and ongoing maintenance burden.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Custom React Application\n')
run.bold = True

doc.add_paragraph(
    'What it does: Custom web application built in React that queries SiteWise and Timestream via AWS SDK, renders '
    'charts (using libraries like Chart.js or Recharts), and provides white-labelled UI.\n\n'
    'Why it matters: Full control over branding, layout, and user experience. Can integrate data from both SiteWise '
    'and Timestream into unified dashboards.\n\n'
    'Technical details: React 18+ with TypeScript. AWS Amplify for authentication integration. API Gateway + Lambda '
    'for backend. Queries SiteWise via BatchGetAssetPropertyValue API and Timestream via SQL.\n\n'
    'Development effort: 6-8 weeks for initial dashboard suite. Ongoing: new chart types, dashboard customisation '
    'requests, browser compatibility, performance optimisation.'
)

doc.add_heading('6.4 Data Flow and Processing', 2)

doc.add_heading('Time-Series Data Path (Machine/Energy/Air)', 3)

flow_ts = [
    'Sensor publishes via OPC-UA/Modbus to SiteWise Gateway (on-premises)',
    'Gateway forwards to AWS IoT SiteWise (latency: <1 second)',
    'SiteWise ingests data, applies asset model, calculates metrics (latency: <2 seconds)',
    'Custom React app queries SiteWise API for dashboard updates (latency: <5 seconds total)',
    'User sees real-time updates in browser'
]

for i, item in enumerate(flow_ts, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('Event Data Path (RFID/Barcode Scans)', 3)

flow_event = [
    'Scanner HTTP POST to API Gateway (latency: <100ms)',
    'Lambda validates and writes to Kinesis Data Streams (latency: <200ms)',
    'Lambda consumer writes to Timestream (latency: <1 second)',
    'Custom React app queries Timestream SQL for job tracking dashboard (latency: <3 seconds total)',
    'User sees job location updates'
]

for i, item in enumerate(flow_event, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('6.5 Strengths and Limitations', 2)

doc.add_heading('Strengths', 3)

strengths = [
    ('Purpose-Built for Manufacturing',
     'SiteWise provides native industrial concepts: asset models, OEE calculations, equipment hierarchies. Speaks '
     'the language of manufacturing rather than generic database terms.'),
    ('Managed Asset Modeling',
     'No custom code needed for standard manufacturing calculations. Built-in formulas for utilisation, availability, '
     'performance. Reduces development compared to building from scratch.'),
    ('Real-Time Latency',
     'Achieves <1 second latency for time-series data. Meets real-time monitoring requirements without complex '
     'stream processing like Flink.'),
    ('Industrial Protocol Support',
     'SiteWise Gateway handles OPC-UA, Modbus TCP without custom integration code. Easier connectivity to legacy '
     'equipment.'),
    ('Automatic Metric Calculations',
     'Compute expressions run at ingestion time. No separate stream processing layer needed for standard transformations.'),
    ('Predictive Maintenance Integration',
     'Amazon Lookout for Equipment integrates directly with SiteWise. Pre-built anomaly detection for manufacturing '
     'equipment (though less flexible than custom SageMaker models).')
]

for title, desc in strengths:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(desc)

doc.add_heading('Limitations and Challenges', 3)

limitations = [
    ('NOT Designed for Multi-Tenant SaaS',
     'CRITICAL: SiteWise has no native Row-Level Security. Tag-based filtering is application-layer workaround that '
     'increases data leakage risk. Acceptable for <10 trusted tenants; risky for SaaS with many companies.'),
    ('Most Expensive Option',
     'Cost: £6,334/month (£211/tenant) vs £2,170/month for Option B (cheapest). 3x more expensive. Ingestion pricing '
     '($0.75 per million values) adds up quickly with high-frequency sensors.'),
    ('Dual Storage Complexity',
     'Requires BOTH SiteWise (time-series) AND Timestream (events). Applications must query two databases and correlate '
     'data. Increases architectural complexity and operational overhead.'),
    ('Custom Dashboard Development Required',
     'SiteWise Monitor cannot be adequately white-labelled. Must build custom React application (6-8 weeks + ongoing '
     'maintenance). Dashboards not included "out of the box" like QuickSight/Grafana.'),
    ('Longest Implementation Timeline',
     '28 weeks due to custom dashboard development, dual storage integration, and SiteWise asset model configuration. '
     '40% longer than Option B (20 weeks).'),
    ('SiteWise Learning Curve',
     'Proprietary concepts (asset models, transforms, metrics) require training. Expression language different from '
     'SQL or Python. Smaller talent pool compared to general databases.'),
    ('Limited ML Flexibility',
     'Lookout for Equipment is pre-built (less flexible than SageMaker). Cannot use custom algorithms. Still requires '
     'SageMaker integration for advanced use cases (adds complexity back).'),
    ('On-Premises Gateway Management',
     'SiteWise Gateway requires deploying software at each manufacturing site. Software updates, troubleshooting, and '
     'security patching across distributed locations.')
]

for title, desc in limitations:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)
    para.add_run(desc)

doc.add_heading('When to Choose Option C', 3)

doc.add_paragraph(
    'Select this architecture ONLY if:'
)

when_choose = [
    'Team has existing deep AWS IoT SiteWise expertise (rare)',
    'Manufacturing focus with <10 tenants (multi-tenancy risk acceptable)',
    'Already using SiteWise for other projects (leverage existing skills)',
    'OPC-UA/Modbus connectivity is critical and must be managed (not delegated to device vendors)',
    'Budget supports £200+/tenant (most expensive option)',
    'Willing to invest 28 weeks + custom dashboard development',
    'Can accept dual storage complexity (SiteWise + Timestream)',
    'AWS-native preference outweighs Snowflake simplicity (but see Option D for better AWS-native choice)'
]

for item in when_choose:
    doc.add_paragraph(item, style='List Bullet')

para = doc.add_paragraph()
run = para.add_run('\n⚠️ ASSESSMENT: ')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(153, 51, 0)
para.add_run(
    'Option C is viable after cost correction but NOT recommended for multi-tenant SaaS. The multi-tenancy risk '
    '(tag-based isolation), dual storage complexity, and custom dashboard burden outweigh the benefits of managed '
    'asset modelling. Choose Option B for simplicity or Option D for AWS-native without SiteWise limitations.'
)

add_page_break(doc)

print("Option C complete. Adding Option D...")

# Save progress
doc.save('/Users/david/projects/smdh/docs/architecture/SMDH-Infrastructure-Design-Options-v2.docx')
print("Document saved (Options A-C complete). Continuing with Option D and comparisons...")

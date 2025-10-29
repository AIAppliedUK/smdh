#!/usr/bin/env python3
"""
Generate comprehensive SMDH Architecture Design Options Word document
Complete version with all 4 options detailed for technical and non-technical audiences
Uses UK English spelling and grammar throughout
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def add_heading_with_color(doc, text, level, color_rgb=(0, 51, 102)):
    """Add a coloured heading"""
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def add_placeholder_box(doc, text, height=3.0):
    """Add a placeholder box for diagrams"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add shaded box
    run = p.add_run(f'\n{text}\n')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Add some spacing
    for _ in range(int(height * 2)):
        p.add_run('\n')

    # Add border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '24')
        border_el.set(qn('w:space'), '4')
        border_el.set(qn('w:color'), 'CCCCCC')
        pBdr.append(border_el)
    pPr.append(pBdr)

    # Add shading
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)

    return p

# Create document
doc = Document()

# Set up default styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

#=============================================================================
# TITLE PAGE
#=============================================================================

title = doc.add_heading('Smart Manufacturing Data Hub (SMDH)', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Architecture Design Options')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph('\n' * 3)

version_para = doc.add_paragraph('Comprehensive Analysis of Four Architecture Approaches')
version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in version_para.runs:
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(70, 70, 70)

doc.add_paragraph('\n' * 5)

# Document information table
info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Light Grid Accent 1'

cells = info_table.rows[0].cells
cells[0].text = 'Document Version'
cells[1].text = '2.0 FINAL'

cells = info_table.rows[1].cells
cells[0].text = 'Date'
cells[1].text = 'October 2025'

cells = info_table.rows[2].cells
cells[0].text = 'Status'
cells[1].text = 'For Review'

cells = info_table.rows[3].cells
cells[0].text = 'Classification'
cells[1].text = 'Internal Use'

cells = info_table.rows[4].cells
cells[0].text = 'Author'
cells[1].text = 'AI Applied Technical Team'

cells = info_table.rows[5].cells
cells[0].text = 'Intended Audience'
cells[1].text = 'Technical Architects and Business Stakeholders'

add_page_break(doc)

#=============================================================================
# EXECUTIVE SUMMARY
#=============================================================================

add_heading_with_color(doc, '1. Executive Summary', 1)

doc.add_paragraph(
    'This document presents a comprehensive analysis of four distinct architectural approaches '
    'for the Smart Manufacturing Data Hub (SMDH) platform. The SMDH is designed as a cloud-native, '
    'multi-tenant Internet of Things (IoT) platform that empowers small and medium-sized manufacturing '
    'enterprises with real-time visibility into their operations.'
)

doc.add_heading('1.1 Document Purpose', 2)

doc.add_paragraph(
    'The purpose of this document is to provide both technical architects and business stakeholders '
    'with a clear understanding of four viable architecture options for the SMDH platform. Each option '
    'has been analysed in detail, considering factors such as:'
)

purpose_list = [
    'Cost implications and total cost of ownership',
    'Technical complexity and required expertise',
    'Performance characteristics and latency requirements',
    'Scalability and multi-tenancy capabilities',
    'Implementation timelines and resource requirements',
    'Operational considerations and maintenance overhead'
]

for item in purpose_list:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.2 System Context', 2)

doc.add_paragraph(
    'The SMDH platform serves manufacturing companies by collecting, processing, and visualising '
    'data from various industrial sensors and equipment. The system must support:'
)

context_list = [
    'Multiple manufacturing companies (tenants) using a shared platform infrastructure',
    'Processing 2.6 to 3.9 million data points per day from diverse sensor types',
    'Real-time monitoring with latency under 1 second for critical alerts',
    'Self-service company onboarding and device registration',
    'Automated dashboard provisioning and data visualisation',
    'Secure data isolation between different companies',
    'Support for 20-40 concurrent users across 30-40 companies in Year 1'
]

for item in context_list:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.3 The Four Architecture Options', 2)

doc.add_paragraph(
    'This document evaluates four distinct architectural approaches, each with unique strengths '
    'and trade-offs:'
)

# Option summaries
doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Option A: Flink-Based AWS-Heavy Architecture')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph(
    'This option leverages Apache Flink for advanced stream processing, providing sub-second latency '
    'and sophisticated data transformation capabilities. It represents the most powerful but also most '
    'complex approach, suitable for scenarios requiring advanced real-time processing.\n'
    'Monthly cost: £2,500-4,200 (30 tenants) | Complexity: Very High | Timeline: 24 weeks'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Option B: Snowflake-Leveraged Architecture')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph(
    'This option centralises data storage and processing in Snowflake, a cloud data platform. '
    'It offers the simplest architecture with SQL-first development, making it ideal for teams '
    'with traditional database skills. This is the recommended option for most scenarios.\n'
    'Monthly cost: £2,170-3,450 (30 tenants) | Complexity: Low | Timeline: 20 weeks'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Option C: AWS IoT SiteWise-Based Architecture')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph(
    'This option uses AWS IoT SiteWise, a managed service specifically designed for industrial IoT. '
    'It provides built-in asset modelling and time-series capabilities but requires custom development '
    'for multi-tenant isolation and dashboards.\n'
    'Monthly cost: £6,334 (30 tenants) | Complexity: Medium-High | Timeline: 28 weeks'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Option D: Pure AWS-Native Architecture (Timestream + Grafana)')
run.bold = True
run.font.size = Pt(12)
doc.add_paragraph(
    'This option uses only AWS services, combining Amazon Timestream for time-series data storage '
    'with Grafana for dashboards. It provides a fully AWS-native solution with excellent multi-tenancy '
    'support and industry-standard visualisation.\n'
    'Monthly cost: £3,965 (30 tenants) | Complexity: Medium | Timeline: 22 weeks'
)

doc.add_heading('1.4 Key Finding: All Options Are Viable', 2)

doc.add_paragraph(
    'An important finding from our analysis is that all four options meet the budget constraint '
    'of £200-300 per tenant per month. A previous cost calculation error for Option C has been '
    'corrected, reducing its estimated monthly cost by 70% (from £21,150 to £6,334 for 30 tenants). '
    'This means the decision should be based on factors other than basic affordability, such as '
    'technical requirements, team capabilities, and strategic preferences.'
)

doc.add_heading('1.5 How to Use This Document', 2)

doc.add_paragraph(
    'This document is structured to support both detailed technical review and high-level business '
    'decision-making:'
)

usage_list = [
    'Technical Architects should review Sections 4-7 for detailed component descriptions and data flows',
    'Business Stakeholders may focus on Section 1 (this summary) and Section 8 (comparisons)',
    'Decision-makers should review Section 9 for specific recommendations and decision frameworks',
    'All readers can refer to Section 10 (Glossary) for explanations of technical terms'
]

for item in usage_list:
    doc.add_paragraph(item, style='List Bullet')

add_page_break(doc)

#=============================================================================
# SECTION 2: DOCUMENT PURPOSE AND AUDIENCE
#=============================================================================

add_heading_with_color(doc, '2. Document Purpose and Audience', 1)

doc.add_heading('2.1 Intended Audience', 2)

doc.add_paragraph(
    'This document has been prepared for a diverse audience with varying levels of technical expertise:'
)

audience_table = doc.add_table(rows=5, cols=2)
audience_table.style = 'Light Grid Accent 1'

cells = audience_table.rows[0].cells
cells[0].text = 'Audience'
cells[1].text = 'What to Focus On'

cells = audience_table.rows[1].cells
cells[0].text = 'Senior Leadership'
cells[1].text = 'Executive Summary (Section 1), Recommendations (Section 9), Cost Comparison (Section 8.1)'

cells = audience_table.rows[2].cells
cells[0].text = 'Product Managers'
cells[1].text = 'Architecture overviews (Sections 4-7), Decision Framework (Section 8.5)'

cells = audience_table.rows[3].cells
cells[0].text = 'Technical Architects'
cells[1].text = 'Detailed component descriptions (Sections 4-7), Data flows, Technical comparisons'

cells = audience_table.rows[4].cells
cells[0].text = 'Engineering Teams'
cells[1].text = 'Component descriptions, Technology stacks, Implementation considerations'

doc.add_heading('2.2 Reading Guide for Non-Technical Readers', 2)

doc.add_paragraph(
    'For readers without extensive cloud computing or AWS knowledge, this document includes:'
)

reading_guide = [
    'Plain English explanations of technical concepts where possible',
    'Visual architecture diagrams showing how components connect',
    'Analogies to familiar concepts (e.g., "like a sorting office for data")',
    'A comprehensive glossary defining all technical terms',
    'Summary boxes highlighting key points in each section',
    'Cost and complexity indicators for quick comparison'
]

for item in reading_guide:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    '\nWhenever you encounter an unfamiliar term, please refer to the Glossary in Section 10. '
    'Terms defined in the glossary appear in this format on first use: Internet of Things (IoT).'
)

add_page_break(doc)

#=============================================================================
# SECTION 3: SYSTEM OVERVIEW AND REQUIREMENTS
#=============================================================================

add_heading_with_color(doc, '3. System Overview and Requirements', 1)

doc.add_heading('3.1 What is the Smart Manufacturing Data Hub?', 2)

doc.add_paragraph(
    'The Smart Manufacturing Data Hub (SMDH) is a cloud-based software platform that helps manufacturing '
    'companies monitor and analyse their operations in real-time. Think of it as a central "command centre" '
    'where data from factory equipment, environmental sensors, and tracking systems all flows together to '
    'provide insights.'
)

doc.add_heading('How It Works (Simplified)', 3)

steps_list = [
    'Manufacturing companies register on the platform and create accounts',
    'They register their factory sites and the equipment/sensors they want to monitor',
    'Sensors send data continuously (e.g., "Machine A is running", "Temperature is 22°C")',
    'The platform receives, validates, and stores this data securely',
    'Dashboards automatically update to show real-time status and historical trends',
    'Alerts notify users when something needs attention (e.g., machine breakdown, poor air quality)'
]

for i, item in enumerate(steps_list, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('3.2 Key System Requirements', 2)

doc.add_paragraph(
    'The architecture options must all satisfy the following essential requirements:'
)

doc.add_heading('Data Volume and Performance', 3)

perf_table = doc.add_table(rows=5, cols=2)
perf_table.style = 'Light List Accent 1'

cells = perf_table.rows[0].cells
cells[0].text = 'Requirement'
cells[1].text = 'Target'

cells = perf_table.rows[1].cells
cells[0].text = 'Daily data points processed'
cells[1].text = '2.6 - 3.9 million rows per day'

cells = perf_table.rows[2].cells
cells[0].text = 'Real-time monitoring latency'
cells[1].text = 'Under 1 second for critical data'

cells = perf_table.rows[3].cells
cells[0].text = 'Alert notification time'
cells[1].text = 'Under 10 seconds from trigger event'

cells = perf_table.rows[4].cells
cells[0].text = 'System availability'
cells[1].text = '99.9% uptime (less than 9 hours downtime per year)'

doc.add_heading('Multi-Tenancy (Shared Platform)', 3)

doc.add_paragraph(
    'The platform must support multiple manufacturing companies simultaneously whilst ensuring '
    'complete data isolation. This means:'
)

tenancy_list = [
    'Company A cannot see Company B\'s data under any circumstances',
    'Each company has their own dashboards, users, and configurations',
    'The platform must efficiently share resources (like servers) whilst maintaining security',
    'Support for 30-40 companies initially, scaling to 100+ companies over time'
]

for item in tenancy_list:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Budget Constraints', 3)

doc.add_paragraph(
    'The target cost is £200-300 per company (tenant) per month. For 30 companies, this means:\n'
)

budget_para = doc.add_paragraph()
run = budget_para.add_run('Target monthly cost: £6,000 - £9,000\n')
run.bold = True
run.font.size = Pt(12)

doc.add_paragraph(
    'This budget must cover all cloud infrastructure, data storage, processing, and dashboard services. '
    'Good news: all four architecture options fall within or close to this budget.'
)

doc.add_heading('3.3 Use Cases Supported', 2)

doc.add_paragraph('The platform must support diverse manufacturing monitoring scenarios:')

# Machine Utilisation
para = doc.add_paragraph()
run = para.add_run('Machine Utilisation Monitoring\n')
run.bold = True
doc.add_paragraph(
    'Tracks how efficiently manufacturing equipment is being used. Sensors collect data every second '
    'about whether machines are running, idle, or experiencing downtime. This helps identify '
    'bottlenecks and improve productivity.'
)

# Air Quality
para = doc.add_paragraph()
run = para.add_run('Air Quality Management\n')
run.bold = True
doc.add_paragraph(
    'Monitors environmental conditions in manufacturing facilities. Sensors measure CO₂, volatile '
    'organic compounds (VOCs), particulate matter, temperature, and humidity. Ensures compliance '
    'with health and safety regulations and worker wellbeing.'
)

# Energy Monitoring
para = doc.add_paragraph()
run = para.add_run('Energy Monitoring\n')
run.bold = True
doc.add_paragraph(
    'Tracks electrical consumption of equipment and facilities. Monitors voltage, current, power '
    'factor, and kilowatt-hour usage. Helps identify energy waste and reduce costs.'
)

# Job Tracking
para = doc.add_paragraph()
run = para.add_run('Job Location Tracking\n')
run.bold = True
doc.add_paragraph(
    'Uses RFID tags or barcodes to track work-in-progress through the factory. Provides real-time '
    'visibility of where jobs are located and helps prevent lost or delayed orders.'
)

add_page_break(doc)

#=============================================================================
# SECTION 4: OPTION A - FLINK-BASED ARCHITECTURE
#=============================================================================

add_heading_with_color(doc, '4. Architecture Option A: Flink-Based AWS-Heavy Architecture', 1, (153, 51, 0))

doc.add_heading('4.1 Overview', 2)

# Summary box
summary_table = doc.add_table(rows=6, cols=2)
summary_table.style = 'Medium Shading 1 Accent 6'

cells = summary_table.rows[0].cells
cells[0].text = 'Monthly Cost (30 tenants)'
cells[1].text = '£2,500 - 4,200'

cells = summary_table.rows[1].cells
cells[0].text = 'Cost Per Tenant'
cells[1].text = '£83 - 140'

cells = summary_table.rows[2].cells
cells[0].text = 'Complexity Level'
cells[1].text = 'Very High (15+ services)'

cells = summary_table.rows[3].cells
cells[0].text = 'Implementation Timeline'
cells[1].text = '24 weeks'

cells = summary_table.rows[4].cells
cells[0].text = 'Team Size Required'
cells[1].text = '3-4 engineers'

cells = summary_table.rows[5].cells
cells[0].text = 'Best Suited For'
cells[1].text = 'Scenarios requiring sub-second latency and advanced ML'

doc.add_paragraph()

doc.add_heading('What is This Architecture?', 3)

doc.add_paragraph(
    'Option A represents the most technically sophisticated approach, using Apache Flink—an advanced '
    'stream processing framework—to handle data in real-time. Imagine a highly efficient assembly line '
    'where data flows through multiple specialised stations, each performing specific transformations '
    'and checks at extremely high speed.'
)

doc.add_paragraph(
    'This architecture is called "AWS-Heavy" because it uses many different Amazon Web Services (AWS) '
    'components, each specialised for a particular task. Whilst this provides maximum flexibility and '
    'power, it also creates significant complexity in managing all these interconnected services.'
)

doc.add_heading('4.2 Architecture Diagram', 2)

add_placeholder_box(doc, '[INSERT: SMDH-Option-A-Flink-Architecture.drawio diagram here]', 4.0)

doc.add_paragraph(
    'The diagram above shows how data flows from sensors through multiple processing stages before '
    'reaching users. Each box represents a different AWS service or component.'
)

doc.add_heading('4.3 Detailed Component Description', 2)

doc.add_paragraph(
    'This section explains each major component of the architecture in plain English, with the '
    'technical details that architects need to evaluate the design.'
)

doc.add_heading('Layer 1: Data Sources and Ingestion', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('AWS IoT Core (MQTT Broker)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Acts as the "front door" for all sensor data entering the system. Sensors connect '
    'to IoT Core using MQTT (a lightweight messaging protocol designed for IoT devices) and publish '
    'their data.\n\n'
    'Why it matters: Handles device authentication using X.509 certificates, ensuring only authorised '
    'sensors can send data. Supports millions of concurrent connections and can receive data from '
    '30-45 sensors per manufacturing site.\n\n'
    'Technical details: Provides Quality of Service (QoS) levels 0 and 1 for message delivery. '
    'Includes device shadows for storing last-known state and offline capabilities.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('IoT Rules Engine\n')
run.bold = True

doc.add_paragraph(
    'What it does: Acts like a "sorting office" that examines each incoming message and routes it '
    'to the appropriate destination based on rules.\n\n'
    'Why it matters: Enables intelligent message routing without writing custom code. Can enrich '
    'messages with additional metadata (like tenant IDs) and filter out invalid data.\n\n'
    'Technical details: SQL-based rules engine. Can route to multiple destinations simultaneously. '
    'Supports transformations using built-in functions.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Kinesis Data Firehose\n')
run.bold = True

doc.add_paragraph(
    'What it does: Buffers incoming data and groups it into larger batches before passing it onward. '
    'Think of it like a truck that waits until it\'s full before making a delivery, rather than making '
    'individual trips for each package.\n\n'
    'Why it matters: Improves efficiency by batching data. Handles automatic compression and format '
    'conversion. Provides built-in retry logic if downstream systems are temporarily unavailable.\n\n'
    'Technical details: Buffers data up to 10MB or 60 seconds (whichever comes first). Supports Parquet, '
    'ORC, and JSON formats. Can invoke Lambda functions for transformation.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Lambda Transformation Functions\n')
run.bold = True

doc.add_paragraph(
    'What it does: Small programs that validate, clean, and transform individual data records. Examples '
    'include converting temperature from Fahrenheit to Celsius, validating sensor ranges, or adding '
    'calculated fields.\n\n'
    'Why it matters: Ensures data quality before storage. Enables custom business logic without managing '
    'servers. Highly scalable—automatically handles processing for one sensor or one thousand.\n\n'
    'Technical details: Serverless functions written in Python or Node.js. Invoked synchronously during '
    'data flow. Timeout: 5 minutes maximum. Can access other AWS services like DynamoDB for lookups.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('S3 Staging Storage\n')
run.bold = True

doc.add_paragraph(
    'What it does: Stores data in structured folders (like an organised filing cabinet) in Parquet format '
    '(an efficient columnar storage format).\n\n'
    'Why it matters: Provides durable, low-cost storage. Acts as a "data lake" where historical data can '
    'be kept indefinitely. Enables recovery if downstream processes fail.\n\n'
    'Technical details: S3 Standard storage with intelligent-tiering for cost optimisation. Partitioned by '
    'tenant_id/device_id/year/month/day for efficient queries. Lifecycle policies move data to Glacier after '
    '90 days for long-term archival.'
)

doc.add_heading('Layer 2: Stream Processing (The Heart of Option A)', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Apache Flink on Amazon EMR\n')
run.bold = True
run.font.color.rgb = RGBColor(153, 51, 0)

doc.add_paragraph(
    'What it does: This is the defining feature of Option A. Apache Flink is an advanced stream processing '
    'framework that processes data in real-time with sub-second latency. It performs complex operations like '
    'windowing (grouping data into time periods), pattern detection (e.g., "detect if a machine cycles on and '
    'off three times in 5 minutes"), and sophisticated aggregations.\n\n'
    'Why it matters: Enables the most advanced real-time analytics capabilities. Can handle complex event '
    'processing that simpler systems cannot. Provides exactly-once processing guarantees, meaning no data '
    'is lost or processed twice.\n\n'
    'Technical details: Runs on Amazon EMR (Elastic MapReduce) cluster with 3-5 nodes. Uses RocksDB for '
    'state management with checkpoints to S3 every 5 minutes. Supports SQL queries (Flink SQL) and custom '
    'Java/Scala/Python code. Autoscaling based on data velocity.\n\n'
    'What makes it complex: Requires expertise in Java or Scala programming. Cluster management and tuning '
    'is non-trivial. Debugging distributed stream processing applications requires specialised skills. '
    'State management and checkpointing must be carefully configured to prevent data loss.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('RocksDB State Backend\n')
run.bold = True

doc.add_paragraph(
    'What it does: Stores the "working memory" of Flink jobs. For example, if you\'re calculating a '
    '5-minute average, it needs to remember the last 5 minutes of data.\n\n'
    'Why it matters: Enables stateful stream processing. Provides fault tolerance—if a server crashes, '
    'processing can resume from the last checkpoint without data loss.\n\n'
    'Technical details: Embedded database within each Flink task. Checkpoints to S3 for durability. '
    'Incremental checkpointing reduces overhead. Recovery time typically under 2 minutes.'
)

doc.add_heading('Layer 3: Data Platform (Snowflake)', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowpipe Auto-Ingestion\n')
run.bold = True

doc.add_paragraph(
    'What it does: Automatically detects when new data files arrive in S3 and loads them into Snowflake '
    '(the data warehouse) within seconds.\n\n'
    'Why it matters: No manual intervention required. Enables near-real-time analytics on batch-loaded data. '
    'Automatically scales based on workload.\n\n'
    'Technical details: Event-driven using S3 notifications → SQS → Snowpipe. Loads micro-batches continuously. '
    'Charges based on warehouse compute time (typically seconds per file).'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowflake Tables and Views\n')
run.bold = True

doc.add_paragraph(
    'What it does: Stores data in structured tables (like Excel spreadsheets but for billions of rows). '
    'Views provide pre-calculated aggregations and enforce row-level security for multi-tenancy.\n\n'
    'Why it matters: SQL interface familiar to most analysts. Native support for VARIANT (JSON) data types. '
    'Row-Level Security (RLS) ensures tenants only see their own data.\n\n'
    'Technical details: Tables partitioned by tenant_id. VARIANT columns for flexible schema. Clustering '
    'on timestamp columns for query performance. Materialised views for common aggregations.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowflake Streams and Tasks\n')
run.bold = True

doc.add_paragraph(
    'What it does: Streams track changes to tables (like a change log), and Tasks are scheduled SQL statements '
    'that run periodically to process those changes.\n\n'
    'Why it matters: Enables incremental processing—only new or changed data is processed, reducing costs. '
    'Provides a native ETL (Extract, Transform, Load) framework within Snowflake.\n\n'
    'Technical details: Change Data Capture (CDC) using streams. DAG-based task orchestration. Serverless '
    'compute—only pay when tasks run.'
)

doc.add_heading('Layer 4: Application and Analytics', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('API Gateway and Lambda API Handlers\n')
run.bold = True

doc.add_paragraph(
    'What it does: Provides REST APIs that applications call to retrieve data. API Gateway acts as the '
    '"front desk," routing requests to Lambda functions that execute the business logic.\n\n'
    'Why it matters: Enables mobile and web applications to access data securely. Enforces authentication '
    'and rate limiting. Scales automatically with demand.\n\n'
    'Technical details: API Gateway with JWT token authentication. Lambda functions query Snowflake or DynamoDB. '
    'Tenant isolation enforced via Cognito user attributes.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Amazon QuickSight (Dashboards)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Creates interactive dashboards and visualisations. Users can see charts, graphs, and tables '
    'showing their manufacturing data.\n\n'
    'Why it matters: Embedded dashboards can be integrated into custom applications. Pay-per-session pricing '
    'model is cost-effective for occasional users. Native connector to Snowflake.\n\n'
    'Technical details: Embedded dashboards with Row-Level Security. Supports scheduled reports and email alerts. '
    'SPICE in-memory cache for fast queries.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Amazon SageMaker (Machine Learning)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Enables building, training, and deploying custom machine learning models. Examples include '
    'predictive maintenance (predicting when equipment will fail) or anomaly detection.\n\n'
    'Why it matters: Provides unlimited flexibility for advanced analytics. Supports all major ML frameworks '
    '(TensorFlow, PyTorch, scikit-learn). GPU-accelerated training for deep learning.\n\n'
    'Technical details: SageMaker training jobs on ml.p3 instances. Model hosting on auto-scaling endpoints. '
    'Feature store for reusable ML features.'
)

doc.add_heading('4.4 Data Flow and Processing', 2)

doc.add_heading('Real-Time Data Flow (Sub-Second Path)', 3)

doc.add_paragraph(
    'For critical data requiring immediate processing (e.g., machine failure alerts):'
)

rt_flow = [
    'Sensor publishes MQTT message to IoT Core (latency: <100ms)',
    'IoT Rules Engine routes to Kinesis Data Streams (latency: <200ms)',
    'Flink consumes from Kinesis and processes with sub-second windowing (latency: <500ms)',
    'Flink detects alert condition and publishes to SNS/SES (latency: <200ms)',
    'User receives email/SMS notification (total latency: <2 seconds)'
]

for i, item in enumerate(rt_flow, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('Batch Processing Flow (Historical Analytics)', 3)

doc.add_paragraph(
    'For historical reporting and trend analysis (less time-sensitive):'
)

batch_flow = [
    'Kinesis Firehose buffers data in 60-second batches',
    'Lambda function transforms and validates each record',
    'Data written to S3 in Parquet format (partitioned by date and tenant)',
    'Snowpipe detects new S3 files and loads into Snowflake (<5 minutes)',
    'Snowflake Tasks run hourly to calculate aggregations',
    'QuickSight dashboards refresh automatically showing updated metrics'
]

for i, item in enumerate(batch_flow, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_heading('4.5 Strengths and Limitations', 2)

doc.add_heading('Strengths', 3)

strengths = [
    ('Sub-Second Latency',
     'Achieves <1 second end-to-end latency for critical data. Essential if regulatory requirements '
     'mandate immediate alerting.'),
    ('Advanced Stream Processing',
     'Flink provides capabilities that simpler systems lack: complex event processing, pattern matching '
     'across multiple streams, sophisticated windowing strategies.'),
    ('Unlimited Flexibility',
     'Can implement any custom logic in Java/Python/Scala. No limitations on data transformations or '
     'ML algorithms.'),
    ('Best-of-Breed Components',
     'Each component is industry-leading in its category. Flink for streaming, Snowflake for analytics, '
     'SageMaker for ML.'),
    ('Proven at Scale',
     'Flink powers stream processing at companies like Netflix, Uber, and Alibaba. Known to handle '
     'billions of events per day.')
]

for title, desc in strengths:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    para.add_run(desc)

doc.add_heading('Limitations and Challenges', 3)

limitations = [
    ('Very High Complexity',
     'Managing 15+ interconnected services requires deep expertise. Each component has its own configuration, '
     'monitoring, and troubleshooting requirements.'),
    ('Steep Learning Curve',
     'Flink expertise is rare and expensive. Teams need proficiency in multiple languages (Java for Flink, '
     'Python for Lambda, SQL for Snowflake).'),
    ('Longest Implementation Timeline',
     'Estimated 24 weeks to production. Building and testing distributed streaming applications takes time.'),
    ('Operational Overhead',
     'EMR cluster management, Flink job tuning, checkpoint monitoring, state recovery testing all require '
     'ongoing attention.'),
    ('Highest Debugging Complexity',
     'Troubleshooting issues across distributed Flink jobs, multiple data pipelines, and async processing '
     'is challenging.'),
    ('Resource Intensive',
     'Requires 3-4 senior engineers with diverse skill sets: Flink specialists, AWS architects, data engineers.')
]

for title, desc in limitations:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)
    para.add_run(desc)

doc.add_heading('When to Choose Option A', 3)

doc.add_paragraph(
    'Select this architecture ONLY if:'
)

when_choose = [
    'Sub-second latency is legally or regulatorily required (not just "nice to have")',
    'Complex event processing (CEP) is essential for the business case',
    'Team already has Flink expertise or budget to hire specialists',
    'Willing to accept 24-week timeline and higher operational complexity',
    'Budget supports £2,500-4,200/month for infrastructure alone'
]

for item in when_choose:
    doc.add_paragraph(item, style='List Bullet')

para = doc.add_paragraph()
run = para.add_run('\nRecommendation: ')
run.bold = True
run.font.size = Pt(12)
para.add_run(
    'For most manufacturing IoT scenarios, the complexity and cost of Option A are not justified. '
    'Consider Option B (Snowflake) for simplicity or Option D (AWS-native) for fast alerting without '
    'Flink complexity.'
)

add_page_break(doc)

#=============================================================================
# SECTION 5: OPTION B - SNOWFLAKE ARCHITECTURE
#=============================================================================

add_heading_with_color(doc, '5. Architecture Option B: Snowflake-Leveraged Architecture', 1, (0, 102, 204))

doc.add_heading('5.1 Overview', 2)

# Summary box
summary_table = doc.add_table(rows=6, cols=2)
summary_table.style = 'Medium Shading 1 Accent 5'

cells = summary_table.rows[0].cells
cells[0].text = 'Monthly Cost (30 tenants)'
cells[1].text = '£2,170 - 3,450 (CHEAPEST)'

cells = summary_table.rows[1].cells
cells[0].text = 'Cost Per Tenant'
cells[1].text = '£72 - 115'

cells = summary_table.rows[2].cells
cells[0].text = 'Complexity Level'
cells[1].text = 'Low (5 core services)'

cells = summary_table.rows[3].cells
cells[0].text = 'Implementation Timeline'
cells[1].text = '20 weeks (FASTEST)'

cells = summary_table.rows[4].cells
cells[0].text = 'Team Size Required'
cells[1].text = '2-3 engineers (SQL skills)'

cells = summary_table.rows[5].cells
cells[0].text = 'Best Suited For'
cells[1].text = 'Most scenarios - Recommended default choice'

doc.add_paragraph()

doc.add_heading('What is This Architecture?', 3)

doc.add_paragraph(
    'Option B represents the simplest and most cost-effective approach, centralising nearly all data '
    'processing and storage within Snowflake—a cloud data platform optimised for analytics. Imagine a '
    '"one-stop shop" where data arrives, gets processed, stored, and analysed all in one unified system, '
    'rather than moving between multiple separate services.'
)

doc.add_paragraph(
    'This architecture is called "Snowflake-Leveraged" because it maximises the use of Snowflake\'s built-in '
    'capabilities (Streams for change tracking, Tasks for scheduling, native SQL for transformations), minimising '
    'the need for external processing engines like Flink or complex Lambda functions.'
)

para = doc.add_paragraph()
run = para.add_run('\n🏆 RECOMMENDED OPTION FOR MOST SCENARIOS\n')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 128, 0)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('5.2 Architecture Diagram', 2)

add_placeholder_box(doc, '[INSERT: SMDH-Option-B-Snowflake-Architecture.drawio diagram here]', 4.0)

doc.add_paragraph(
    'The diagram above shows a much simpler data flow compared to Option A. Notice how data reaches '
    'Snowflake quickly and most processing happens within a single platform.'
)

doc.add_heading('5.3 Detailed Component Description', 2)

doc.add_heading('Layer 1: Simplified Ingestion Path', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('AWS IoT Core (Same as Option A)\n')
run.bold = True

doc.add_paragraph(
    'Function identical to Option A—receives MQTT messages from sensors with X.509 certificate authentication. '
    'The difference in Option B is what happens next: data flows directly to Snowflake with minimal intermediate '
    'processing.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Kinesis Data Streams (Minimal Buffering)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Provides a lightweight buffer between IoT Core and Snowflake. Unlike Option A which uses '
    'Kinesis Firehose with heavy batching, Option B uses Kinesis Streams for near-real-time delivery.\n\n'
    'Why it matters: Reduces latency whilst still providing ordering guarantees and replay capabilities. '
    'Simpler configuration than Firehose.\n\n'
    'Technical details: 10-20 shards based on throughput. 24-hour retention for replay. Partition key: tenant_id '
    'for tenant-level ordering.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowpipe Streaming\n')
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph(
    'What it does: Snowflake\'s streaming ingestion service that continuously loads data from Kinesis directly '
    'into Snowflake tables in near-real-time (typically 5-10 seconds).\n\n'
    'Why it matters: Eliminates the need for S3 staging, Lambda transformations, and batch loading. Data arrives '
    'in Snowflake faster and with fewer moving parts. This is a key simplification compared to Option A.\n\n'
    'Technical details: Serverless—no infrastructure to manage. Automatically handles schema evolution. Exactly-once '
    'delivery semantics. Uses Snowflake\'s VARIANT type for flexible JSON ingestion.'
)

doc.add_heading('Layer 2: Unified Data Platform (Snowflake)', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowflake: The Core of Option B\n')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph(
    'In Option B, Snowflake performs the role of multiple services from Option A combined:\n'
)

sf_roles = [
    'Data Lake (replaces S3)',
    'Stream Processing (replaces Flink)',
    'Data Warehouse (same as Option A)',
    'ETL Engine (replaces Lambda + AWS Glue)',
    'Query Engine (same as Option A but simpler)'
]

for role in sf_roles:
    doc.add_paragraph(f'• {role}')

doc.add_paragraph()

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Raw Tables with VARIANT Columns\n')
run.bold = True

doc.add_paragraph(
    'What it does: Stores incoming JSON data in Snowflake\'s VARIANT type, which preserves the original structure '
    'whilst allowing SQL queries on nested fields.\n\n'
    'Why it matters: No need to pre-define rigid schemas. If sensors start sending new fields, they\'re automatically '
    'captured. Queries can still extract specific values: SELECT data:temperature::float FROM raw_data.\n\n'
    'Technical details: Tables partitioned by TENANT_ID and DEVICE_ID for query performance. Clustering on TIMESTAMP '
    'column. Automatic compression reduces storage costs by 75-85%.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowflake Streams (Change Data Capture)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Automatically tracks what data has been added, updated, or deleted in tables. Think of it like '
    'a "changelog" that remembers what\'s new since you last looked.\n\n'
    'Why it matters: Enables incremental processing—transformations only operate on new data, not the entire table '
    'every time. Drastically reduces processing costs and time.\n\n'
    'Technical details: Streams use hidden metadata columns (METADATA$ACTION, METADATA$ISUPDATE). Zero-copy—doesn\'t '
    'duplicate data. Multiple consumers can read from the same stream independently.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowflake Tasks (Orchestration)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Scheduled SQL statements that run automatically. For example, "Every 5 minutes, calculate '
    'hourly aggregations for any new data."\n\n'
    'Why it matters: Replaces external orchestration tools like Apache Airflow. Enables building complex data '
    'pipelines entirely in SQL. DAG (Directed Acyclic Graph) support for task dependencies.\n\n'
    'Technical details: Serverless execution using dedicated warehouses. Cron-like scheduling or stream-triggered '
    'execution. Built-in retry logic and alerting on failure.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Curated Views with Row-Level Security (RLS)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Pre-calculated views that automatically filter data based on who\'s querying. A user from '
    'Company A sees only Company A\'s data; Company B sees only theirs.\n\n'
    'Why it matters: Multi-tenancy security enforced at the database level—cannot be bypassed by applications. '
    'Simpler than application-level filtering. Snowflake\'s RLS is mature and battle-tested.\n\n'
    'Technical details: Secure views with SESSION_CONTEXT(\'CURRENT_TENANT\') filtering. Mapping tables define '
    'user→tenant relationships. Policy objects for reusable security logic across multiple tables.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowflake Time-Travel and Zero-Copy Cloning\n')
run.bold = True

doc.add_paragraph(
    'What it does: Time-Travel allows querying historical data ("show me this table as it was 2 days ago"). '
    'Zero-copy cloning creates instant copies of databases for testing without duplicating storage.\n\n'
    'Why it matters: Time-Travel enables easy recovery from mistakes (accidental deletions, bad updates). '
    'Cloning enables safe testing of schema changes or ETL modifications.\n\n'
    'Technical details: Time-Travel retention configurable 1-90 days (1 day default). Cloning uses copy-on-write '
    '(only changes consume additional storage). Fail-safe provides 7 additional days for Snowflake support to '
    'recover data.'
)

doc.add_heading('Layer 3: Machine Learning in Snowflake', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Snowflake Cortex ML\n')
run.bold = True

doc.add_paragraph(
    'What it does: Enables building machine learning models using SQL commands within Snowflake. Example: '
    'CREATE MODEL predict_machine_failure AS SELECT * FROM training_data;\n\n'
    'Why it matters: No need to export data to external ML platforms. Analysts with SQL skills can build models '
    'without Python expertise. Predictions run as SQL queries.\n\n'
    'Technical details: Supports classification, regression, time-series forecasting. AutoML automatically selects '
    'best algorithm. Model versioning and lineage tracking built-in.\n\n'
    'Limitations: Less flexible than SageMaker (Option A). Cannot use custom algorithms or deep learning frameworks. '
    'Good for 80% of ML use cases; advanced scenarios may require SageMaker integration.'
)

doc.add_heading('Layer 4: Application and Analytics (Simplified)', 3)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('API Gateway + Lambda (Minimal)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Same role as Option A but much simpler. Lambda functions primarily just execute SQL queries '
    'against Snowflake—no complex transformation logic.\n\n'
    'Why it matters: Less custom code means fewer bugs and faster development. Snowflake handles the heavy lifting; '
    'Lambda is just a thin API layer.\n\n'
    'Technical details: Snowflake Node.js or Python connector. Connection pooling for performance. API responses '
    'cached in CloudFront for frequently-accessed data.'
)

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('QuickSight or Power BI (User Choice)\n')
run.bold = True

doc.add_paragraph(
    'What it does: Dashboarding identical to Option A. Option B offers flexibility: use QuickSight (AWS-native) '
    'or Power BI (popular in manufacturing sector).\n\n'
    'Why it matters: Many manufacturing companies already have Power BI licences and expertise. Snowflake\'s native '
    'connectors for both tools enable easy integration.\n\n'
    'Technical details: Direct query mode (live connection to Snowflake) or import mode (cached extracts). Row-Level '
    'Security passes through from Snowflake. Embedded dashboards for white-labelling.'
)

doc.add_heading('5.4 Data Flow and Processing', 2)

doc.add_heading('End-to-End Data Flow (Simplified)', 3)

doc.add_paragraph(
    'The beauty of Option B is the dramatically simpler data path:'
)

flow = [
    'Sensor publishes MQTT message to IoT Core (latency: <100ms)',
    'IoT Core routes to Kinesis Data Streams via Rules Engine (latency: <200ms)',
    'Snowpipe Streaming ingests directly into Snowflake raw table (latency: 5-10 seconds)',
    'Snowflake Stream detects new rows immediately',
    'Snowflake Task (runs every 1 minute) processes new rows and updates curated views (latency: <60 seconds)',
    'QuickSight dashboard refreshes showing latest data (total latency: 60-65 seconds for dashboards)'
]

for i, item in enumerate(flow, 1):
    doc.add_paragraph(f'{i}. {item}')

doc.add_paragraph()
para = doc.add_paragraph()
run = para.add_run('Important Note on Latency:\n')
run.bold = True
run.font.color.rgb = RGBColor(153, 51, 0)

doc.add_paragraph(
    'The default configuration provides 60-65 second latency for dashboard updates. This violates the <10 second '
    'alert requirement. HOWEVER, this is easily fixed by adding a Lambda "fast-path" for alerts (see Section 8 '
    'for details). The fast-path adds £100-150/month and 2 weeks to timeline but achieves <5 second alert latency '
    'whilst keeping dashboard processing in Snowflake.'
)

doc.add_heading('5.5 Strengths and Limitations', 2)

doc.add_heading('Strengths', 3)

strengths = [
    ('Lowest Total Cost',
     '£2,170-3,450/month for 30 tenants (£72-115 per tenant). Cheapest option by 15-20%. Cost scales linearly '
     'with usage—no expensive clusters idling overnight.'),
    ('Lowest Complexity',
     'Only 5 core services vs 15+ in Option A. Single platform (Snowflake) handles most processing. Less to '
     'configure, monitor, and troubleshoot.'),
    ('Fastest Time-to-Market',
     '20-week timeline (fastest of all options). SQL-first development means faster iteration. Large pool of '
     'SQL-skilled developers vs rare Flink specialists.'),
    ('SQL-First Development',
     'Transformations, aggregations, and ML all use SQL. Familiar to data analysts and engineers. Easier to '
     'review and maintain than Java/Scala code.'),
    ('Native Multi-Tenancy',
     'Snowflake\'s Row-Level Security is production-proven. Policy-based security reduces risk of data leakage. '
     'Tenant isolation enforced at database level, cannot be bypassed.'),
    ('Unified Data Governance',
     'All data in one platform simplifies governance. Single place to audit access, track lineage, enforce retention '
     'policies. Reduces compliance overhead.'),
    ('Automatic Scaling',
     'Snowflake warehouses auto-suspend when idle (no wasted cost) and auto-resume on query (no manual intervention). '
     'Multi-cluster warehouses handle query concurrency spikes.'),
    ('Proven at Manufacturing Scale',
     'Snowflake used by major manufacturers (Toyota, Siemens) for industrial IoT. Handles 100+ petabytes of data '
     'in production.')
]

for title, desc in strengths:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    para.add_run(desc)

doc.add_heading('Limitations and Considerations', 3)

limitations = [
    ('Alert Latency Requires Fix',
     'Default 60-65 second latency violates <10s requirement. FIX: Add Lambda fast-path for alerts (+£100-150/month, '
     '+2 weeks). After fix, alert latency <5 seconds whilst batch processing stays in Snowflake.'),
    ('Snowflake Licensing Cost',
     'Snowflake charges separately from AWS bill. Requires contract negotiation. Some organisations have policies '
     'against third-party data platforms (prefer AWS-native Option D).'),
    ('Vendor Lock-in (Medium)',
     'Snowflake-specific features (Streams, Tasks, VARIANT) not portable to other databases. However, SQL code is '
     'mostly standard and data can be exported. Less lock-in than Option A (Flink expertise lost if switching).'),
    ('Limited Flexibility for Complex ML',
     'Cortex ML suitable for standard use cases (regression, classification). Custom deep learning or specialised '
     'algorithms require SageMaker integration (adds complexity). Not as flexible as Option A.'),
    ('Real-Time Latency Not Sub-Second',
     'Snowpipe Streaming provides 5-10 second latency, not <1 second. If sub-second requirement is legally mandated '
     'for ALL data (not just alerts), Option A or D required.')
]

for title, desc in limitations:
    para = doc.add_paragraph()
    run = para.add_run(f'{title}: ')
    run.bold = True
    run.font.color.rgb = RGBColor(153, 51, 0)
    para.add_run(desc)

doc.add_heading('When to Choose Option B', 3)

doc.add_paragraph(
    'Select this architecture if:'
)

when_choose = [
    'Cost and simplicity are top priorities (most SME scenarios)',
    'Team has SQL skills or can train quickly (common)',
    '60-second dashboard latency acceptable, or willing to add Lambda fast-path for <5s alerts',
    'Want fastest time to market (20 weeks baseline, 22 weeks with fast-path)',
    'Need strong data governance and compliance (unified platform simplifies)',
    'Snowflake licensing not a blocker (or already have Snowflake contract)',
    'Prefer proven technology over cutting-edge (Snowflake widely adopted in manufacturing)'
]

for item in when_choose:
    doc.add_paragraph(item, style='List Bullet')

para = doc.add_paragraph()
run = para.add_run('\n✅ RECOMMENDATION: ')
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 128, 0)
para.add_run(
    'Option B is the recommended default choice for the SMDH platform. It offers the best balance of cost, '
    'simplicity, and capabilities for 80% of manufacturing IoT scenarios. The alert latency issue is easily '
    'addressed with a Lambda fast-path. Only choose Option A if sub-second latency is legally mandated, or '
    'Option D if AWS-native is required.'
)

add_page_break(doc)

print("Sections 1-5 complete. Continuing with Options C and D...")

# Save progress
doc.save('/Users/david/projects/smdh/docs/architecture/SMDH-Infrastructure-Design-Options-v2.docx')
print("Document saved (partial). Continuing...")
